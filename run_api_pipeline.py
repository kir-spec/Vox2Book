#!/usr/bin/env python3
"""
Vox2Book — Direct Neural API Pipeline
Pure Python API Client for OpenAI, DeepSeek, Anthropic Claude, LM Studio, and Ollama.
Processes input texts/transcripts from inputs/raw_texts/ and outputs DOCX manuscripts to output/books/.
"""

import os
import re
import json
import sys
import urllib.request
import urllib.error

# Ensure stdout handles UTF-8 on Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

SYSTEM_PROMPT = """Ты — главный редактор литературного издательства.
Твоя задача — взять сырую устную расшифровку аудиозаписи (Whisper STT transcript) и превратить её в вычитанный, чистый и грамотный литературный текст.

ПРАВИЛА:
1. Разбей сплошной монолог на логические абзацы и четкие предложения с правильной пунктуацией.
2. Исправь все ошибочные технические термины распознавания устной речи:
   - "ссд" / "ssd" -> SSD-накопитель
   - "те из бы" / "юсб" -> USB
   - "а дата" -> ADATA
   - "35 размер" -> 3.5-дюймовый
   - "планктом" -> планктон
   - "в стране джетал" -> Western Digital
   - "трансценд" -> Transcend
3. Удали слова-паразиты ("ну", "э-э", "как бы", "значит").
4. Сохрани авторский смысл, но сделай стиль литературным.
5. Верни ТОЛЬКО вычитанный готовый текст книги без вводных фраз."""

def apply_typography(text: str) -> str:
    """Applies Russian publishing typography rules."""
    if not text:
        return ""
    text = re.sub(r'"([^"]+)"', r'«\1»', text)
    text = re.sub(r' - ', ' — ', text)
    text = re.sub(r'\b(из|из за)\b', 'из-за', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(из под)\b', 'из-под', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(все таки|всё таки)\b', 'всё-таки', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(в обшем|вобщем)\b', 'в общем', text, flags=re.IGNORECASE)
    return text

def call_neural_api(prompt_text: str, config: dict) -> str:
    """Calls selected Neural API provider."""
    provider = config.get("api_provider", "openai").lower()
    api_key = config.get("api_key", "").strip() or os.environ.get("OPENAI_API_KEY", "")
    model = config.get("model", "")

    if provider in ["openai", "deepseek", "lmstudio"]:
        if provider == "deepseek":
            url = "https://api.deepseek.com/chat/completions"
            model_name = model or "deepseek-chat"
        elif provider == "lmstudio":
            base_url = config.get("lmstudio_url", "http://localhost:1234").rstrip("/")
            url = f"{base_url}/v1/chat/completions"
            model_name = model or "local-model"
        else:
            url = "https://api.openai.com/v1/chat/completions"
            model_name = model or "gpt-4o-mini"

        if provider != "lmstudio" and (not api_key or "YOUR_API_KEY" in api_key):
            print(f"[Notice] No valid API key configured for {provider}. Using rule-based proofreading engine.")
            return prompt_text

        headers = {"Content-Type": "application/json"}
        if provider != "lmstudio" and api_key:
            headers["Authorization"] = f"Bearer {api_key}"

        payload = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt_text}
            ],
            "temperature": 0.3
        }

        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers)
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"[Notice] Neural API ({provider}): {e}. Using rule-based proofreading engine.")
            return prompt_text

    elif provider in ["anthropic", "claude"]:
        if not api_key or "YOUR_API_KEY" in api_key:
            print("[Notice] No valid API key configured for Anthropic. Using rule-based proofreading engine.")
            return prompt_text
            
        url = "https://api.anthropic.com/v1/messages"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01"
        }
        payload = {
            "model": model or "claude-3-5-sonnet-20240620",
            "max_tokens": 4096,
            "system": SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": prompt_text}]
        }
        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers)
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data["content"][0]["text"].strip()
        except Exception as e:
            print(f"[Notice] Anthropic API: {e}. Using rule-based proofreading engine.")
            return prompt_text

    elif provider == "ollama":
        ollama_url = config.get("ollama_url", "http://localhost:11434").rstrip("/")
        url = f"{ollama_url}/api/generate"
        payload = {
            "model": model or "llama3",
            "prompt": f"{SYSTEM_PROMPT}\n\nИсходный текст:\n{prompt_text}",
            "stream": False
        }
        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers={"Content-Type": "application/json"})
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data.get("response", "").strip()
        except Exception as e:
            print(f"[Notice] Ollama: {e}. Using rule-based proofreading engine.")
            return prompt_text

    return prompt_text

def build_docx_manuscript(text: str, output_path: str, title: str = ""):
    """Builds a formatted Word DOCX manuscript."""
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    if title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)
        run.font.bold = True
        doc.add_paragraph()

    # Process Body Paragraphs
    paragraphs = text.split("\n\n")
    for para_str in paragraphs:
        para_clean = para_str.strip()
        if not para_clean:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

        # Format dialogues
        formatted_text = apply_typography(para_clean)
        if formatted_text.startswith("-") or formatted_text.startswith("—"):
            formatted_text = "— " + formatted_text.lstrip("-—").strip()

        run = p.add_run(formatted_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[SUCCESS] Manuscript saved to: '{output_path}'")

def main():
    print("=========================================================")
    print("   Vox2Book — Direct Neural API Publishing Pipeline   ")
    print("=========================================================")

    # Ensure required folder structure exists
    os.makedirs("inputs/raw_texts", exist_ok=True)
    os.makedirs("inputs/audio", exist_ok=True)
    os.makedirs("output/books", exist_ok=True)

    config_path = "config.json"
    if not os.path.exists(config_path):
        default_config = {
            "api_provider": "openai",
            "api_key": os.environ.get("OPENAI_API_KEY", ""),
            "model": "gpt-4o-mini",
            "ollama_url": "http://localhost:11434",
            "lmstudio_url": "http://localhost:1234",
            "genre": "prose",
            "title": "",
            "subtitle": ""
        }
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(default_config, f, indent=2, ensure_ascii=False)

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    # Determine input file path
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        raw_files = [os.path.join("inputs/raw_texts", f) for f in os.listdir("inputs/raw_texts") if f.endswith(('.txt', '.md', '.docx'))]
        if raw_files:
            input_file = raw_files[0]
        else:
            input_file = "inputs/raw_texts/sample.txt"
            with open(input_file, "w", encoding="utf-8") as f:
                f.write("Поместите ваш текст сюда.")

    # ALWAYS output processed file into output/books/ directory!
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_file = os.path.join("output/books", f"{base_name}.docx")

    print(f"Input:  '{input_file}'")
    print(f"Output: '{output_file}' (Saved in output/books/)")
    print(f"Neural Provider: {config.get('api_provider')} ({config.get('model')})")

    # Read input content with multi-encoding support (UTF-8, Windows-1251)
    raw_bytes = open(input_file, "rb").read()
    if raw_bytes.startswith(b'\xef\xbb\xbf'):
        raw_bytes = raw_bytes[3:]

    try:
        raw_text = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            raw_text = raw_bytes.decode("cp1251")
        except UnicodeDecodeError:
            raw_text = raw_bytes.decode("utf-8", errors="ignore")

    print("Sending text to Neural API for proofreading...")
    edited_text = call_neural_api(raw_text, config)

    print("Formatting and building DOCX manuscript...")
    build_docx_manuscript(edited_text, output_file, config.get("title", ""))

if __name__ == "__main__":
    main()
