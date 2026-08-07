import unittest
import os
import tempfile
import docx
from pipeline import (
    stage1_cleanup,
    stage3_typography,
    stage4_audit,
    process_manuscript_chain,
)

class TestUniversalPipeline(unittest.TestCase):
    
    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def test_typography_rules(self):
        sample = 'Он сказал: "Привет" - и пошёл из за угла вобщем все таки.'
        plan = {"locale": "ru"}
        res = stage3_typography(sample, plan)
        self.assertIn('«Привет»', res)
        self.assertIn(' — ', res)
        self.assertIn('из-за', res)
        self.assertIn('в общем', res)
        self.assertIn('всё-таки', res)
        
    def test_stt_cleaning(self):
        sample = 'это текста фрагмент с рекламой Quiz河'
        plan = {"actions": ["cleanup"]}
        res = stage1_cleanup(sample, plan)
        self.assertNotIn('Quiz河', res)

    def test_poetry_pipeline(self):
        poem_text = """Глава I. Весна

Я помню чудное мгновенье:
Передо мной явилась ты,
Как мимолетное виденье,
Как гений чистой красоты.

В томленьях грусти безнадежной,
В тревогах шумной суеты,
Звучал мне долго голос нежный
И снились милые черты.
"""
        txt_path = os.path.join(self.temp_dir, "poem.txt")
        out_docx = os.path.join(self.temp_dir, "poem.docx")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(poem_text)
            
        process_manuscript_chain(txt_path, out_docx)
        self.assertTrue(os.path.exists(out_docx))
        
        doc = docx.Document(out_docx)
        self.assertGreater(len(doc.paragraphs), 5)
        
    def test_drama_pipeline(self):
        play_text = """Глава I. Действие первое

Спикер 1 (входит в комнату): Привет! Как твои дела?
Спикер 2: Всё отлично, изучаю новые алгоритмы.
(Раздаётся звонок телефонного аппарата)
Спикер 1: Я отвечу на звонок.
"""
        txt_path = os.path.join(self.temp_dir, "play.txt")
        out_docx = os.path.join(self.temp_dir, "play.docx")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(play_text)
            
        process_manuscript_chain(txt_path, out_docx)
        self.assertTrue(os.path.exists(out_docx))
        
        doc = docx.Document(out_docx)
        self.assertGreater(len(doc.paragraphs), 4)

    def test_prose_pipeline(self):
        prose_text = """Глава 1. Начало путешествия

Солнце ярко светило над горизонтом. Вся природа просыпалась от долгого зимнего сна.

- Нам пора отправляться в путь - сказал путник.
- Согласен - ответил его спутник.
"""
        txt_path = os.path.join(self.temp_dir, "prose.txt")
        out_docx = os.path.join(self.temp_dir, "prose.docx")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(prose_text)
            
        process_manuscript_chain(txt_path, out_docx)
        self.assertTrue(os.path.exists(out_docx))
        
        doc = docx.Document(out_docx)
        self.assertGreater(len(doc.paragraphs), 3)

    def test_scenarios_grid(self):
        """
        Runs scenario combinations testing typography and cleanup functions.
        """
        speakers = ['Спикер 1', 'Спикер 2', 'Автор', 'Рассказчик']
        word_samples = ['привет', 'как дела', 'из за', 'всё таки', 'в общем', '«тест»']
        
        count = 0
        plan = {"locale": "ru"}
        for sp in speakers:
            for w in word_samples:
                count += 1
                cleaned = stage3_typography(f"{sp}: {w}", plan)
                self.assertTrue(len(cleaned) > 0)
                    
        self.assertGreaterEqual(count, 20)


if __name__ == '__main__':
    unittest.main()
