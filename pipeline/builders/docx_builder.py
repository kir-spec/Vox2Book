import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import MONTHS_RU, GENRE_PRESETS, DEFAULT_SPEAKER_COLORS

def build_docx_document(elements, output_path, genre="prose", title=None, subtitle=None):
    """
    Universal DOCX document generator supporting prose, poetry, drama, dialogue, and academic.
    Dynamically maps speaker colors for multi-speaker dialogues.
    """
    doc = docx.Document()
    preset = GENRE_PRESETS.get(genre, GENRE_PRESETS['prose'])
    
    font_name = preset.get('font_name', 'Times New Roman')
    font_size_pt = preset.get('font_size_pt', 12)
    line_spacing = preset.get('line_spacing', 1.2)
    
    # 1. Document Title
    doc_title = title if title else preset.get('title', 'Литературный макет')
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run(doc_title)
    r_title.bold = True
    r_title.font.name = font_name
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    # 2. Subtitle
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(12)
        r_sub = p_sub.add_run(subtitle)
        r_sub.italic = True
        r_sub.font.name = font_name
        r_sub.font.size = Pt(13)
        r_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        
    # Speaker color registry
    speaker_color_map = {}
    color_palette = DEFAULT_SPEAKER_COLORS
    
    # 3. Process Elements
    current_date_key = None
    
    for item in elements:
        itype = item.get('type')
        body = item.get('edited_body', item.get('body', '')).strip()
        speaker = item.get('speaker', 'Speaker')
        
        # Dynamically assign speaker color
        if speaker and speaker not in speaker_color_map:
            color_idx = len(speaker_color_map) % len(color_palette)
            speaker_color_map[speaker] = RGBColor(*color_palette[color_idx])
            
        if genre == 'dialogue' and item.get('type') == 'message':
            dt = item.get('dt')
            if dt:
                date_key = (dt.year, dt.month, dt.day)
                if date_key != current_date_key:
                    current_date_key = date_key
                    month_name = MONTHS_RU.get(dt.month, '')
                    date_str = f"📅 {dt.day} {month_name} {dt.year} г."
                    
                    p_date = doc.add_paragraph()
                    p_date.paragraph_format.space_before = Pt(16)
                    p_date.paragraph_format.space_after = Pt(6)
                    r_date = p_date.add_run(date_str)
                    r_date.bold = True
                    r_date.font.name = "Calibri"
                    r_date.font.size = Pt(14)
                    r_date.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                    
            m_type = "Голосовое" if item.get('msg_type') == 'voice' else "Текст"
            time_str = item.get('time_str', '00:00')
            
            p_msg = doc.add_paragraph()
            p_msg.paragraph_format.space_before = Pt(2)
            p_msg.paragraph_format.space_after = Pt(5)
            p_msg.paragraph_format.line_spacing = 1.15
            
            r_sp = p_msg.add_run(speaker)
            r_sp.bold = True
            r_sp.font.name = "Calibri"
            r_sp.font.size = Pt(11)
            r_sp.font.color.rgb = speaker_color_map.get(speaker, RGBColor(0x1F, 0x49, 0x7D))
            
            r_meta = p_msg.add_run(f" [{time_str}] [{m_type}]: ")
            r_meta.bold = True
            r_meta.font.name = "Calibri"
            r_meta.font.size = Pt(10)
            r_meta.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
            
            r_b = p_msg.add_run(body)
            r_b.font.name = "Calibri"
            r_b.font.size = Pt(11)
            r_b.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
            continue
            
        if itype == 'heading':
            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_head.paragraph_format.space_before = Pt(14)
            p_head.paragraph_format.space_after = Pt(6)
            r_h = p_head.add_run(body)
            r_h.bold = True
            r_h.font.name = font_name
            r_h.font.size = Pt(16)
            r_h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            continue
            
        if genre == 'poetry':
            if itype == 'stanza_break':
                p_sb = doc.add_paragraph()
                p_sb.paragraph_format.space_before = Pt(0)
                p_sb.paragraph_format.space_after = Pt(10)
                continue
                
            p_v = doc.add_paragraph()
            p_v.paragraph_format.space_before = Pt(0)
            p_v.paragraph_format.space_after = Pt(1)
            p_v.paragraph_format.left_indent = Pt(36)
            r_v = p_v.add_run(body)
            r_v.font.name = "Georgia"
            r_v.font.size = Pt(11.5)
            continue
            
        if genre == 'drama':
            if itype == 'character_name':
                p_c = doc.add_paragraph()
                p_c.paragraph_format.space_before = Pt(8)
                p_c.paragraph_format.space_after = Pt(2)
                r_c = p_c.add_run(body)
                r_c.bold = True
                r_c.font.name = "Arial"
                r_c.font.size = Pt(11)
                r_c.font.color.rgb = speaker_color_map.get(speaker, RGBColor(0x1F, 0x49, 0x7D))
            elif itype == 'stage_direction':
                p_sd = doc.add_paragraph()
                p_sd.paragraph_format.space_before = Pt(2)
                p_sd.paragraph_format.space_after = Pt(4)
                r_sd = p_sd.add_run(body)
                r_sd.italic = True
                r_sd.font.name = "Arial"
                r_sd.font.size = Pt(10.5)
                r_sd.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            else:
                p_d = doc.add_paragraph()
                p_d.paragraph_format.space_before = Pt(0)
                p_d.paragraph_format.space_after = Pt(4)
                r_d = p_d.add_run(body)
                r_d.font.name = "Arial"
                r_d.font.size = Pt(11)
            continue
            
        # Default Prose / Academic / Essay
        if body:
            p_p = doc.add_paragraph()
            p_p.paragraph_format.space_before = Pt(0)
            p_p.paragraph_format.space_after = Pt(6)
            p_p.paragraph_format.line_spacing = line_spacing
            if genre == 'prose' and not body.startswith('—'):
                p_p.paragraph_format.first_line_indent = Pt(18)
            r_p = p_p.add_run(body)
            r_p.font.name = font_name
            r_p.font.size = Pt(font_size_pt)
            
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Universal DOCX document saved: {output_path}")
