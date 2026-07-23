import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import MONTHS_RU

def create_book_docx(messages, output_path, subtitle="Полная хроника общения (2026 г.)\nСобеседники: Kir и Амфи"):
    """
    Creates DOCX book document matching exact formatting of sample docx Голосовые_сообщения_2024-2025.docx.
    """
    doc = docx.Document()
    
    # 1. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run("Диалоги и голосовые расшифровки")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    # 2. Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run(subtitle)
    r_sub.italic = True
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    # 3. Empty separator
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(0)
    
    current_date_key = None
    
    for msg in messages:
        text = msg.get('edited_body', '').strip()
        if not text:
            continue
            
        dt = msg.get('dt')
        if dt:
            date_key = (dt.year, dt.month, dt.day)
            if date_key != current_date_key:
                current_date_key = date_key
                month_name = MONTHS_RU.get(dt.month, '')
                date_str = f"📅 {dt.day} {month_name} {dt.year} г."
                
                # Date header paragraph
                p_date = doc.add_paragraph()
                p_date.paragraph_format.space_before = Pt(16)
                p_date.paragraph_format.space_after = Pt(6)
                r_date = p_date.add_run(date_str)
                r_date.bold = True
                r_date.font.name = "Calibri"
                r_date.font.size = Pt(14)
                r_date.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                
        speaker = msg.get('speaker', 'Kir')
        m_type = "Голосовое" if msg.get('type') == 'voice' else "Текст"
        time_str = msg.get('time_str', '00:00')
        
        # Message paragraph
        p_msg = doc.add_paragraph()
        p_msg.paragraph_format.space_before = Pt(2)
        p_msg.paragraph_format.space_after = Pt(5)
        p_msg.paragraph_format.line_spacing = 1.15
        
        # Run 1: Speaker Name
        r_speaker = p_msg.add_run(speaker)
        r_speaker.bold = True
        r_speaker.font.name = "Calibri"
        r_speaker.font.size = Pt(11)
        if speaker == "Kir":
            r_speaker.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        else:  # Амфи
            r_speaker.font.color.rgb = RGBColor(0x80, 0x00, 0x40)
            
        # Run 2: Metadata Prefix
        meta_text = f" [{time_str}] [{m_type}]: "
        r_meta = p_msg.add_run(meta_text)
        r_meta.bold = True
        r_meta.font.name = "Calibri"
        r_meta.font.size = Pt(10)
        r_meta.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        
        # Run 3: Body Text
        r_body = p_msg.add_run(text)
        r_body.font.name = "Calibri"
        r_body.font.size = Pt(11)
        r_body.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
        
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document DOCX saved successfully to: {output_path}")
