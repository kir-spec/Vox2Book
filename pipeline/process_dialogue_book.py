#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Vox2Book — dialogue chronicle processor (month headers, literary STT edit, DOCX)."""

from __future__ import annotations

import argparse
import os
import re
import sys

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# Import pipeline stages from project root
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
from pipeline import stage1_stt_cleanup, stage3_publisher_typography, stage4_quality_auditor  # noqa: E402

MONTH_GENITIVE = {
    "января": "Январь",
    "февраля": "Февраль",
    "марта": "Март",
    "апреля": "Апрель",
    "мая": "Май",
    "июня": "Июнь",
    "июля": "Июль",
    "августа": "Август",
    "сентября": "Сентябрь",
    "октября": "Октябрь",
    "ноября": "Ноябрь",
    "декабря": "Декабрь",
}

DAY_RE = re.compile(r"^📅\s*(\d{1,2}\s+\S+\s+\d{4}\s*г\.?)\s*$")
DAY_RE_PLAIN = re.compile(r"^(\d{1,2}\s+\S+\s+\d{4}\s*г\.?)\s*$")
MSG_RE = re.compile(
    r"^(Амфи|Анфи|Kir)\s+\[(\d{1,2}:\d{2})\]\s+\[(Голосовое|Текст)\]:\s*(.*)$",
    re.DOTALL,
)

BRAND_FIXES = [
    (r"\bbluetooth\b", "Bluetooth"),
    (r"\bgoogle\b", "Google"),
    (r"\bsony\b", "Sony"),
    (r"\biphone\b", "iPhone"),
    (r"\btelegram\b", "Telegram"),
    (r"\btiktok\b", "TikTok"),
    (r"\byoutube\b", "YouTube"),
    (r"\bwindows\b", "Windows"),
    (r"\bwifi\b", "Wi-Fi"),
    (r"\bwi-fi\b", "Wi-Fi"),
    (r"\busb\b", "USB"),
    (r"\bssd\b", "SSD"),
    (r"\btype-c\b", "Type-C"),
    (r"\btype c\b", "Type-C"),
]

STT_PHRASE_FIXES = [
    (r"\bне жны\b", "не нужны"),
    (r"\bпол раз\b", "пару раз"),
    (r"\bвоно\b", "всё"),
    (r"\bохереть\b", "охренеть"),
    (r"\bохеревать\b", "охреневать"),
    (r"\bПривет а\b", "Привет! А"),
    (r"\bпривет а\b", "Привет! А"),
    (r"\bдва две\b", "две"),
    (r"\bфирмы Сони\b", "фирмы Sony"),
    (r"\bколонка фирмы Сони\b", "колонка фирмы Sony"),
    (
        r"Привет!?\s*А,?\s*что это вообще за вещь такая\s+"
        r"(?:я\s+)?(?:расставлю|рассматриваю)[,\s]*температур[ауеы]?[.\s]*"
        r"(?:короче\s+)?(?:валясь|валяюсь)[,\s]*",
        "Привет! А что это вообще за вещь такая? У меня температура, я валяюсь. ",
    ),
    (
        r"что это я не понимаю две дырочки провод ты",
        "Не понимаю: две дырочки, провод — ты",
    ),
    (r"провод — ты, хотя бы", "провод — ты хотя бы"),
]

CLAUSE_STARTERS = [
    "короче говоря",
    "поэтому",
    "значит",
    "то есть",
    "в общем",
    "слушай",
    "кстати",
    "ладно",
    "смотри",
    "понимаешь",
    "короче",
    "потому что",
    "так что",
    "а если",
    "и вот",
    "и так",
    "вот",
    "ну",
]

SUBORDINATE = [
    "потому что",
    "то есть",
    "так что",
    "чтобы",
    "когда",
    "если",
    "хотя",
    "который",
    "которая",
    "которые",
    "которое",
    "что",
]


def normalize_speaker(name: str) -> str:
    if name in ("Амфи", "Анфи"):
        return "Анфи"
    return name


def replace_amfi(text: str) -> str:
    return text.replace("Амфи", "Анфи")


def extract_month_key(day_line: str) -> str | None:
    m = re.search(r"\d{1,2}\s+(\S+)\s+(\d{4})", day_line)
    if not m:
        return None
    gen, year = m.group(1).lower(), m.group(2)
    month = MONTH_GENITIVE.get(gen)
    if not month:
        return None
    return f"{month} {year}"


def literary_edit_body(body: str, *, is_voice: bool) -> str:
    if not body:
        return body

    text = body.strip()
    text = replace_amfi(text)
    text = stage1_stt_cleanup(text)

    for pattern, repl in BRAND_FIXES + STT_PHRASE_FIXES:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)

    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"([,;:])(?!\s)", r"\1 ", text)
    text = re.sub(r"\s+([,.;:!?…])", r"\1", text)
    text = re.sub(r"\.{2,}", "…", text)

    if is_voice:
        for conj in SUBORDINATE:
            text = re.sub(
                rf"([а-яёa-z0-9»\"])\s+({conj})\s",
                rf"\1, \2 ",
                text,
                flags=re.IGNORECASE,
            )

        for starter in sorted(CLAUSE_STARTERS, key=len, reverse=True):
            text = re.sub(
                rf"([.!?…])\s*({starter})\s",
                lambda m, s=starter: f"{m.group(1)} {s.capitalize()} ",
                text,
                flags=re.IGNORECASE,
            )
            text = re.sub(
                rf"([а-яёa-z0-9»\"])\s+({starter})\s",
                lambda m, s=starter: f"{m.group(1)}. {s.capitalize()} ",
                text,
                flags=re.IGNORECASE,
            )

        text = re.sub(r"([.!?…])\s*([а-яё])", lambda m: m.group(1) + " " + m.group(2).upper(), text)

    if text and text[0].islower():
        text = text[0].upper() + text[1:]

    text = re.sub(r"\s+([,.;:!?…])", r"\1", text)
    text = re.sub(r"([.!?…])([А-ЯA-Z])", r"\1 \2", text)
    text = re.sub(
        r"([а-яё])(\s)(Короче|Значит|Поэтому|Слушай|Кстати|В общем)\b",
        r"\1. \3",
        text,
    )

    if text and text[-1] not in ".!?…:»\"":
        text += "."

    text = re.sub(r"\bА, что\b", "А что", text)
    text = re.sub(r"\bа, что\b", "а что", text)
    text = re.sub(r"\bты, хотя бы\b", "ты хотя бы", text)

    text = stage3_publisher_typography(text)
    return text


def parse_day_line(line: str) -> str | None:
    line = line.strip()
    m = DAY_RE.match(line) or DAY_RE_PLAIN.match(line)
    return m.group(1).strip() if m else None


def process_paragraphs(paragraphs: list[str], *, subtitle: str) -> list[tuple[str, str]]:
    """Return list of (kind, text) where kind: title|subtitle|month|day|message."""
    out: list[tuple[str, str]] = []
    current_month: str | None = None

    out.append(("title", "Диалоги Анфи и Kir"))
    out.append(("subtitle", subtitle))

    for raw in paragraphs:
        line = replace_amfi(raw.strip())
        if not line:
            continue

        if line in ("Диалоги и устная речь", "Диалоги Анфи и Kir"):
            continue
        if line.startswith("Полная хроника") or line.startswith("Собеседники:"):
            continue

        day = parse_day_line(line)
        if day:
            mk = extract_month_key(day)
            if mk and mk != current_month:
                current_month = mk
                out.append(("month", mk))
            out.append(("day", day))
            continue

        m = MSG_RE.match(line)
        if m:
            speaker = normalize_speaker(m.group(1))
            time_s = m.group(2)
            kind = m.group(3)
            body = literary_edit_body(m.group(4), is_voice=(kind == "Голосовое"))
            out.append(
                (
                    "message",
                    f"{speaker} [{time_s}] [{kind}]: {body}",
                )
            )
            continue

        # Fallback: keep unknown lines with amfi fix
        if line.startswith("📅"):
            day = parse_day_line(line)
            if day:
                mk = extract_month_key(day)
                if mk and mk != current_month:
                    current_month = mk
                    out.append(("month", mk))
                out.append(("day", day))
        else:
            out.append(("other", line))

    return out


def add_styled_paragraph(doc: Document, kind: str, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"

    if kind == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(24)
        run.bold = True
        p.paragraph_format.space_after = Pt(12)
        return

    if kind == "subtitle":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(18)
        return

    if kind == "month":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(18)
        run.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        return

    if kind == "day":
        run.font.size = Pt(12)
        run.italic = True
        run.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Inches(0)
        return

    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def build_dialogue_docx(blocks: list[tuple[str, str]], output_path: str) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for kind, text in blocks:
        add_styled_paragraph(doc, kind, text)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def read_docx_paragraphs(path: str) -> list[str]:
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]


def fix_amfi_in_docx_inplace(path: str) -> int:
    doc = Document(path)
    count = 0
    for p in doc.paragraphs:
        if "Амфи" in p.text:
            for run in p.runs:
                if "Амфи" in run.text:
                    run.text = run.text.replace("Амфи", "Анфи")
                    count += 1
            if "Амфи" in p.text:
                # fallback if runs split oddly
                inline = p.text.replace("Амфи", "Анфи")
                if inline != p.text:
                    p.text = inline
                    count += 1
    doc.save(path)
    return count


def audit_blocks(blocks: list[tuple[str, str]]) -> list[str]:
    messages = [t for k, t in blocks if k == "message"]
    joined = "\n\n".join(messages)
    return stage4_quality_auditor(joined)


def process_file(
    input_path: str,
    output_path: str,
    *,
    subtitle: str,
    fix_source: bool = False,
) -> dict:
    if fix_source:
        fix_amfi_in_docx_inplace(input_path)

    paras = read_docx_paragraphs(input_path)
    blocks = process_paragraphs(paras, subtitle=subtitle)
    issues = audit_blocks(blocks)
    build_dialogue_docx(blocks, output_path)

    msg_count = sum(1 for k, _ in blocks if k == "message")
    month_count = sum(1 for k, _ in blocks if k == "month")
    return {
        "messages": msg_count,
        "months": month_count,
        "audit_issues": len(issues),
        "output": output_path,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Process dialogue DOCX for Vox2Book")
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--subtitle", required=True)
    parser.add_argument("--fix-source", action="store_true")
    args = parser.parse_args()

    stats = process_file(
        args.input,
        args.output,
        subtitle=args.subtitle,
        fix_source=args.fix_source,
    )
    print(f"OK: {stats['output']}")
    print(f"  messages={stats['messages']} months={stats['months']} audit_warnings={stats['audit_issues']}")


if __name__ == "__main__":
    main()
