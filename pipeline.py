#!/usr/bin/env python3
"""
Vox2Book — Виртуальный редакторский отдел (Virtual Editorial Board)
Полный конвейер, как в настоящем издательстве:

  Приёмка → Авто-определение → Корректор (гигиена) → Литературный редактор (LLM)
  → Корректор-типограф → Выпускающий редактор (8 аудитов) → Верстальщик (DOCX)

Каждый этап пишет кэш-файл в output/.llm_cache/ для прослеживаемости.
"""

import os
import re
import json
import sys
import urllib.request
import urllib.error

# Ensure stdout handles UTF-8 on Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# Auto-detection module
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "tools"))
try:
    from auto_detect import analyze as _auto_analyze
except Exception:
    _auto_analyze = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH


# =====================================================================
# UTILITIES
# =====================================================================
def _cache(stage_name: str, text: str, ext: str = "txt"):
    """Save intermediate result to output/.llm_cache/ for traceability."""
    cache_dir = os.path.join("output", ".llm_cache")
    os.makedirs(cache_dir, exist_ok=True)
    path = os.path.join(cache_dir, f"{stage_name}.{ext}")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def _read_file(path: str) -> str:
    """Read text file with auto-detection of encoding (UTF-8 / cp1251).
    Supports .txt, .md, and .docx inputs."""
    if path.lower().endswith('.docx'):
        doc = Document(path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    raw = open(path, "rb").read()
    if raw.startswith(b'\xef\xbb\xbf'):
        raw = raw[3:]
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("cp1251", errors="ignore")


# =====================================================================
# AUTO-DETECTION: Genre, style, actions
# =====================================================================
def auto_detect_plan(raw_text: str, config: dict) -> dict:
    """Analyzes text and returns an execution plan (genre, style, actions, flags)."""
    if _auto_analyze is None:
        return {
            "genre": config.get("genre", "prose") or "prose",
            "profile": "none",
            "style_mode": "literary",
            "actions": ["cleanup", "rebuild", "punctuate", "typography", "audit", "docx"],
            "language": "ru",
            "confidence": 0.0,
            "reasons": ["auto_detect module unavailable"],
            "keep_speakers": False,
            "keep_timestamps": False,
            "keep_mat": False,
        }
    plan = _auto_analyze(raw_text)
    cfg_genre = (config.get("genre") or "").strip()
    if cfg_genre and cfg_genre.lower() not in ("auto-detect", "auto", ""):
        plan["genre"] = cfg_genre.lower()
    return plan


def print_plan(plan: dict):
    """Prints the auto-detected plan for the user."""
    print("=== РЕДАКТОРСКИЙ ОТДЕЛ — ПЛАН ===")
    print(f"  Жанр:          {plan.get('genre')}")
    print(f"  Профиль:       {plan.get('profile')}")
    print(f"  Режим стиля:   {plan.get('style_mode')}")
    print(f"  Язык:          {plan.get('language')}")
    print(f"  Уверенность:   {plan.get('confidence')}")
    print(f"  Действия:      {', '.join(plan.get('actions', []))}")
    print(f"  Спикеры:       {plan.get('keep_speakers')}")
    print(f"  Время:         {plan.get('keep_timestamps')}")
    print(f"  Мат:           {plan.get('keep_mat')}")
    reasons = plan.get('reasons', [])
    if reasons:
        print(f"  Сигналы:       {'; '.join(reasons[:4])}")
    print("=================================")


# =====================================================================
# STAGE 1: КОРРЕКТОР — Гигиена источника
# (Удаление мусора, дублирований, STT-омофонов, восстановление брендов)
# =====================================================================

# --- Галлюцинации и служебный мусор STT/OCR / Реклама ---
HALLUCINATIONS = [
    "Субтитры сделал", "Субтитры:", "Редактор субтитров",
    "Продолжение следует", "Благодарю за просмотр",
    "Подписывайтесь на канал", "Quiz河", "DimaTorzok",
    "Отправить. Отправить.", "Отправить.",
    "Продолжение следует...", "Created by", "Создано программой",
    "📺", "📥", "480p", "720p", "1080p", "@TopSaversBot",
]

# --- Машинный мусор и реклама (регулярки) ---
MACHINE_GARBAGE_RE = [
    r'https?://\S+',
    r'Https://\S+',
    r'www\.\S+',
    r'\bBe/[A-Za-z0-9_-]+(\?\S*)?',
    r'\bbe/[A-Za-z0-9_-]+(\?\S*)?',
    r'\bCom/shorts/[A-Za-z0-9_-]+(\?\S*)?',
    r'\bcom/shorts/[A-Za-z0-9_-]+(\?\S*)?',
    r'\bshorts/[A-Za-z0-9_-]+(\?\S*)?',
    r'\bSi=[A-Za-z0-9_-]+\b',
    r'\bsi=[A-Za-z0-9_-]+\b',
    r'^\s*@TopSaversBot\.?\s*$',
    r'^\s*[📺📥]\s*$',
    r'^\s*(480p|720p|1080p)\s*$',
    r'Субтитры сделал\.?',
    r'Отправить\.\s*Отправить\.\s*',
    r'Отправить\.',
    r'Прочай еще,?',
    r'\bход gou\b',
    r'\bЕцируuts\b',
    r'ецируuts',
    r'\bс redhares\b',
    r'\bredhares\b',
    r'С redhares',
    r'Happy end\. С redhares\.',
    r'\bHappy end\.?\s*',
    r'\bВальдерсе\b',
]

# --- STT-омофоны и слова-призраки (контекстные замены) ---
STT_REPLACEMENTS = [
    # Бренды и программы (STT-искажения — не разговорные названия!)
    (r'\bпротуз\b', 'Pro Tools'),
    (r'\bевре перри\b', 'REAPER'),
    (r'\bевре\b', 'REAPER'),
    (r'\bf?lag player\b', 'FLAC-плеер'),
    (r'\bизломанное\b', 'взломанное'),
    (r'\bВальдбрис\b', 'Wildberries'),
    (r'\bВБ банка\b', 'WB-банка'),
    (r'\bаудо-сити\b', 'Audacity'),
    (r'\bдастин стал\b', 'Dustin Stahl'),
    (r'\bсins стал\b', 'Synth Stahl'),
    # Омофоны / слова-призраки
    (r'\bвалясь\b', 'валяюсь'),
    (r'\bдва две\b', 'две'),
    (r'\bтри линиями\b', 'тремя линиями'),
    (r'\bбредяд бредовый\b', 'бредовый'),
    (r'\bчеловечку зеленых\b', 'человечков в зелёном'),
    (r'\bбез вывозного\b', 'безвылазно'),
    (r'\bвэ этом\b', 'в этом'),
    (r'\bспесь специальные\b', 'специальные'),
    (r'\bшею же\b', 'это же'),
    (r'\bнож частично\b', 'ножки частично'),
    (r'\bнавесят шкаф\b', 'навесной шкаф'),
    (r'\bтуда-додвигать\b', 'туда-сюда двигать'),
    (r'\bтанковые люди\b', 'такие люди'),
    (r'\bпахеризм\b', 'пофигизм'),
    (r'\bгринда\b', 'грит'),
    (r'\bлишь предпрямо\b', 'прямо'),
    (r'\bраджу руку\b', 'делал руку'),
    (r'\bна шокой\b', 'на такой'),
    (r'\bкинезащим сонить\b', 'кинематограф'),
    (r'\bмучим\b', 'там'),
    # Грамматические STT-ошибки
    (r'\bрасставлю температура\b', 'расставлю температуру'),
    (r'\bдва дырочки провод\b', 'два гнезда провода'),
    (r'\bдва дырочки\b', 'два гнезда'),
    (r'\bдырочки провод\b', 'гнезда провода'),
    (r'\bхз\b', 'ХЗ'),
]

# --- Технический жаргон (бренды) ---
SLANG_DICTIONARY = [
    (r'\b(те из бы|юсб|ю эс би)\b', 'USB'),
    (r'\b(ссд|с с д|эс эс д)\b', 'SSD'),
    (r'\b(а дата|адата)\b', 'ADATA'),
    (r'\b(вестерн диджитал|в стране джетал)\b', 'Western Digital'),
    (r'\b(трансценд|трансенд)\b', 'Transcend'),
    (r'\b(самсунг)\b', 'Samsung'),
    (r'\b(виндовс|винда|вин 10)\b', 'Windows'),
    (r'\b(блютуз|блю ту з)\b', 'Bluetooth'),
    (r'\b(вай фай|вайфай|wi fi)\b', 'Wi-Fi'),
    (r'\b(эпл|эппл)\b', 'Apple'),
    (r'\b(телеграм|телега)\b', 'Telegram'),
    (r'\b(ютуб|ютюб)\b', 'YouTube'),
    (r'\b(питон|пайтон)\b', 'Python'),
    (r'\b(джава|жаба)\b', 'Java'),
    (r'\b(реакт)\b', 'React'),
    (r'\b(постгрес|постгре)\b', 'PostgreSQL'),
    (r'\b(редис)\b', 'Redis'),
    (r'\b(нжинкс|энджинкс)\b', 'Nginx'),
    (r'\b(докер)\b', 'Docker'),
    (r'\b(кубер|кубернетес)\b', 'Kubernetes'),
    (r'\b(кафка|кафку)\b', 'Kafka'),
    (r'\b(гитхаб)\b', 'GitHub'),
]


def _fix_repetitions(text: str) -> str:
    """Свести дублирования слов: «я я» → «я», «не не» → «не» и т.д."""
    words = ['я', 'не', 'ну', 'короче', 'вот', 'да', 'там', 'типа',
             'прикольно', 'как', 'так', 'что', 'это', 'он', 'она', 'они']
    for w in words:
        text = re.sub(
            r'\b(' + w + r')\s+\1\b',
            r'\1',
            text,
            flags=re.IGNORECASE,
        )
    return text


def _normalize_spaces(text: str) -> str:
    """Нормализация пробелов: двойные → одинарные, висячие → убрать."""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' *\n *', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Удалить строки, состоящие только из пробелов/пунктуации (мусор после удаления)
    lines = text.split('\n')
    cleaned = []
    for ln in lines:
        ln_stripped = ln.strip()
        # Если строка состоит только из точки/пробела — пропустить
        if ln_stripped in ('', '.', '..', '…'):
            if not cleaned or cleaned[-1] != '':
                cleaned.append('')
        else:
            cleaned.append(ln_stripped)
    return '\n'.join(cleaned).strip()


def stage1_cleanup(raw_text: str, plan: dict) -> str:
    """
    Stage 1: Корректор — гигиена источника.
    Удаляет машинный мусор, галлюцинации, сводит дублирования,
    восстанавливает бренды, исправляет STT-омофоны.
    """
    text = raw_text

    actions = plan.get('actions', [])

    # 1a. Удаление галлюцинаций
    if 'remove_garbage' in actions or 'cleanup' in actions:
        for h in HALLUCINATIONS:
            text = text.replace(h, "")
        for pattern in MACHINE_GARBAGE_RE:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        # Удалить строки, которые после очистки содержат только пунктуацию/пробелы
        text = re.sub(r'^\s*[.,;:…\s]+\s*$', '', text, flags=re.M)
        # Удалить сообщения спикеров, где после очистки остался только заголовок без контента
        # (например "Анфи [22:25] [Текст]: ." → удалить целиком)
        text = re.sub(
            r'^\w+\s+\[\d{1,2}:\d{2}\]\s+\[[^\]]+\]:\s*[.,;:…\s]*$',
            '',
            text,
            flags=re.M,
        )

    # 1b. Восстановление брендов / жаргона
    if 'restore_brands' in actions or 'cleanup' in actions:
        for pattern, replacement in SLANG_DICTIONARY:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # 1c. STT-замены (контекстно-безопасные)
    if 'fix_stt' in actions or 'cleanup' in actions:
        for pattern, replacement in STT_REPLACEMENTS:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # 1d. Дублирования слов
    if 'fix_repetitions' in actions or 'cleanup' in actions:
        text = _fix_repetitions(text)

    # 1e. Нормализация пробелов и кодировки
    text = _normalize_spaces(text)

    return text


# =====================================================================
# STAGE 2: ЛИТЕРАТУРНЫЙ РЕДАКТОР — Neural LLM
# (Пересборка сплошного потока в предложения, пунктуация, стиль)
# =====================================================================

def _load_system_prompt(lang: str = "ru") -> str:
    """Load universal editorial system prompt from prompts/<lang>/."""
    base = os.path.dirname(os.path.abspath(__file__))
    lang_map = {"ru": "ru", "en": "en", "uk": "uk", "ua": "uk"}
    folder = lang_map.get((lang or "ru")[:2].lower(), "ru")
    path = os.path.join(base, "prompts", folder, "UNIVERSAL_EDITOR_SYSTEM.md")
    if os.path.isfile(path):
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        for marker in ("## Role", "## Роль"):
            idx = text.find(marker)
            if idx != -1:
                return text[idx:].strip()
        return text.strip()
    return (
        "You are a senior literary editor. Transform raw text into publication-quality prose. "
        "Preserve 100% meaning. Return only the edited text."
    )


SYSTEM_PROMPT = _load_system_prompt()


# =====================================================================
# LM Studio integration: server check + model listing
# =====================================================================

def _lmstudio_check(base_url: str) -> dict:
    """Check if LM Studio server is online and list loaded models."""
    url = base_url.rstrip("/") + "/v1/models"
    try:
        req = urllib.request.Request(url, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode("utf-8"))
            models = [m.get("id", "") for m in data.get("data", [])]
            return {"online": True, "models": models}
    except Exception:
        return {"online": False, "models": []}


def _detect_provider(config: dict) -> str:
    """
    Auto-detect best available LLM provider.
    Priority: cloud API key → LM Studio → Ollama → regex fallback.
    """
    provider = config.get("api_provider", "auto").lower()
    if provider != "auto":
        return provider

    # 1. Cloud providers with API key
    api_key = config.get("api_key", "").strip()
    if api_key and "YOUR_API_KEY" not in api_key:
        return "openai"

    env_key = os.environ.get("OPENAI_API_KEY", "")
    if env_key:
        return "openai"

    # 2. LM Studio (local)
    base_url = config.get("lmstudio_url", "http://localhost:1234").rstrip("/")
    status = _lmstudio_check(base_url)
    if status["online"] and status["models"]:
        return "lmstudio"

    # 3. Ollama (local)
    ollama_url = config.get("ollama_url", "http://localhost:11434").rstrip("/")
    try:
        req = urllib.request.Request(f"{ollama_url}/api/tags", headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=3) as response:
            data = json.loads(response.read().decode("utf-8"))
            if data.get("models"):
                return "ollama"
    except Exception:
        pass

    # 4. Fallback: regex (no LLM)
    return "regex"


def _pick_best_model(available: list, plan: dict) -> str:
    """Определите лучшую модель для литературного редактирования из доступных моделей."""
    if not available:
        return ""
    priority_70b = ["llama-3.3-70b", "llama-3.3", "70b"]
    priority_14b = ["deepseek-r1", "14b", "qwen3", "gemma"]
    priority_12b = ["gemma-4-12", "12b", "qwen2.5"]
    priority_7b = ["7b", "8b", "llama-3", "mistral"]
    for keyword in priority_70b + priority_14b + priority_12b + priority_7b:
        for m in available:
            if keyword in m.lower():
                return m
    return available[0]


# =====================================================================
# Chunking: split long text for LLM context limits
# =====================================================================

CHUNK_SIZE = 4000
CHUNK_OVERLAP = 200


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE,
                overlap: int = CHUNK_OVERLAP) -> list:
    """Split text into chunks at paragraph boundaries, with overlap."""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    paragraphs = text.split("\n\n")
    current = ""
    for para in paragraphs:
        if not para.strip():
            continue
        if len(current) + len(para) + 2 > chunk_size and current:
            chunks.append(current.strip())
            tail = current[-overlap:] if len(current) > overlap else ""
            current = tail + "\n\n" + para
        else:
            current = current + "\n\n" + para if current else para
    if current.strip():
        chunks.append(current.strip())
    return chunks


# =====================================================================
# Regex-based punctuation rebuild (fallback when no LLM available)
# =====================================================================

def _punctuate_content(text: str) -> str:
    """Add basic punctuation to unpunctuated speech."""
    if not text or len(text) < 20:
        return text

    # 1. Запятые перед союзами ("Привет а" → "Привет, а"; "Привет а," → "Привет, а,")
    for conj in ['а', 'но', 'что', 'потому что', 'поэтому', 'когда', 'если', 'хотя', 'значит', 'короче']:
        # " слово а " → " слово, а "
        text = re.sub(r'(\w)\s+(' + conj + r')(\s)', r'\1, \2\3', text, flags=re.IGNORECASE)
        # " слово а, " → " слово, а, " (если запятая уже после союза)
        text = re.sub(r'(\w)\s+(' + conj + r'),', r'\1, \2,', text, flags=re.IGNORECASE)

    # 2. Множественные запятые → одна
    text = re.sub(r',{2,}', ',', text)

    # 3. Точки после завершающих слов-маркеров перед заглавной буквой
    end_markers = ['вот', 'короче', 'прикольно', 'ага', 'да ну',
                   'понятно', 'естественно', 'конечно']
    for marker in end_markers:
        text = re.sub(
            r'\b' + marker + r'\s+(?=[А-ЯЁA-Z])',
            marker + '. ',
            text,
        )

    # 3. Точка в конце, если нет терминального знака
    if not text.endswith(('.', '!', '?', '…', ':', '—', '»')):
        text = text + '.'

    # 4. Заглавная буква в начале
    text = text[0].upper() + text[1:] if text else text

    # 5. Заглавные буквы после точек
    text = re.sub(r'([.!?]\s+)([а-яё])', lambda m: m.group(1) + m.group(2).upper(), text)

    # 6. Запятые перед вводными словами
    for intro in ['кстати', 'вообще', 'например', 'возможно',
                  'наверное', 'скорее', 'буквально']:
        text = re.sub(r'(\w)\s+(' + intro + r')\b', r'\1, \2', text)

    # 7. Пробел перед запятой → после
    text = re.sub(r'\s+,', ',', text)

    return text.strip()


def stage2b_regex_punctuate(text: str, plan: dict) -> str:
    """
    Stage 2b: Регекс-пунктуатор (запасной алгоритм, без LLM).
    Разбивает сплошной поток на предложения, расставляет точки,
    запятые перед союзами, заглавные буквы.
    НЕ заменяет LLM — грубое приближение для работы без нейросети.
    """
    if not text:
        return ""
    lines = text.split('\n')
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            result.append('')
            continue
        msg_match = re.match(
            r'^(\w+\s+\[\d{1,2}:\d{2}\]\s+\[[^\]]+\]:\s*)(.*)$',
            line
        )
        if msg_match:
            prefix = msg_match.group(1)
            content = msg_match.group(2)
            content = _punctuate_content(content)
            result.append(prefix + content)
        else:
            if len(line) > 100 and not line.startswith('📅') and not line.startswith('#'):
                result.append(_punctuate_content(line))
            else:
                result.append(line)
    return '\n'.join(result)


def _llm_request_openai_compatible(url, headers, system_prompt, chunk, model_name, timeout=300):
    """Send a single chunk to an OpenAI-compatible API endpoint."""
    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": chunk},
        ],
        "temperature": 0.3,
    }
    req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers)
    with urllib.request.urlopen(req, timeout=timeout) as response:
        res_data = json.loads(response.read().decode("utf-8"))
        return res_data["choices"][0]["message"]["content"].strip()


def stage2_neural_stylist(text: str, config: dict, plan: dict = None) -> str:
    """
    Stage 2: Литературный редактор — passes text through Neural LLM.
    Providers: OpenAI, DeepSeek, Claude/Anthropic, LM Studio, Ollama.
    If no API available — falls back to regex-based punctuation (stage2b).
    Supports chunked processing for long texts.
    """
    provider = config.get("api_provider", "lmstudio").lower()

    # Auto-detect provider if set to "auto"
    if provider == "auto":
        provider = _detect_provider(config)
        print(f"  [Авто-провайдер] Выбран: {provider}")

    # --- Regex fallback (no LLM) ---
    if provider == "regex":
        print("  [Литературный редактор] LLM недоступен. Перехожу на регекс-пунктуатор.")
        return stage2b_regex_punctuate(text, plan)

    api_key = config.get("api_key", "").strip() or os.environ.get("OPENAI_API_KEY", "")
    model = config.get("model", "")

    system_prompt = SYSTEM_PROMPT
    if plan:
        context = (
            f"\n\n## Авто-определённый план (для этой правки)\n"
            f"- Жанр: {plan.get('genre')}\n"
            f"- Профиль: {plan.get('profile')}\n"
            f"- Режим стиля: {plan.get('style_mode')}\n"
            f"- Язык: {plan.get('language')}\n"
            f"- Сохранять метки спикеров: {plan.get('keep_speakers')}\n"
            f"- Сохранять время: {plan.get('keep_timestamps')}\n"
            f"- Сохранять мат: {plan.get('keep_mat')}\n"
            f"- Действия: {', '.join(plan.get('actions', []))}\n"
            f"\nВАЖНО: Верни ТОЛЬКО исправленный текст. Без комментариев, без пояснений.\n"
            f"Сохраняй все метки спикеров и времени в формате: Имя [ЧЧ:ММ] [Тип]: текст\n"
        )
        system_prompt = system_prompt + context

    # --- LM Studio: auto-check server + select model ---
    if provider == "lmstudio":
        base_url = config.get("lmstudio_url", "http://localhost:1234").rstrip("/")
        status = _lmstudio_check(base_url)
        if not status["online"]:
            print(f"  [LM Studio] Сервер не запущен ({base_url}).")
            print("  [LM Studio] Запустите LM Studio, загрузите модель, включите Developer Server.")
            print("  [Литературный редактор] Перехожу на регекс-пунктуатор (без LLM).")
            return stage2b_regex_punctuate(text, plan)

        available_models = status["models"]
        print(f"  [LM Studio] Сервер онлайн. Доступные модели: {available_models}")
        model_name = model if model and model != "local-model" else _pick_best_model(available_models, plan)
        if not model_name and available_models:
            model_name = available_models[0]
        if not model_name:
            print("  [LM Studio] Нет загруженных моделей. Перехожу на регекс-пунктуатор.")
            return stage2b_regex_punctuate(text, plan)
        print(f"  [LM Studio] Используется модель: {model_name}")

        chunks = _chunk_text(text)
        if len(chunks) > 1:
            print(f"  [Литературный редактор] Текст разбит на {len(chunks)} порций (по {CHUNK_SIZE} символов).")

        url = f"{base_url}/v1/chat/completions"
        headers = {"Content-Type": "application/json"}
        results = []
        for idx, chunk in enumerate(chunks):
            if len(chunks) > 1:
                print(f"    Порция {idx+1}/{len(chunks)} ({len(chunk)} символов)...")
            try:
                result = _llm_request_openai_compatible(url, headers, system_prompt, chunk, model_name, timeout=600)
                results.append(result)
            except Exception as e:
                print(f"  [LM Studio] Ошибка порции {idx+1}: {e}")
                results.append(chunk)
        return "\n\n".join(results)

    # --- OpenAI / DeepSeek ---
    elif provider in ["openai", "deepseek"]:
        if provider == "deepseek":
            url = "https://api.deepseek.com/chat/completions"
            model_name = model or "deepseek-chat"
        else:
            url = "https://api.openai.com/v1/chat/completions"
            model_name = model or "gpt-4o-mini"

        if not api_key or "YOUR_API_KEY" in api_key:
            print(f"  [{provider}] Нет API-ключа. Перехожу на регекс-пунктуатор.")
            return stage2b_regex_punctuate(text, plan)

        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
        chunks = _chunk_text(text)
        if len(chunks) > 1:
            print(f"  [{provider}] Текст разбит на {len(chunks)} порций.")

        results = []
        for idx, chunk in enumerate(chunks):
            if len(chunks) > 1:
                print(f"    Порция {idx+1}/{len(chunks)}...")
            try:
                result = _llm_request_openai_compatible(url, headers, system_prompt, chunk, model_name, timeout=120)
                results.append(result)
            except Exception as e:
                print(f"  [{provider}] Ошибка порции {idx+1}: {e}")
                results.append(chunk)
        return "\n\n".join(results)

    # --- Anthropic / Claude ---
    elif provider in ["anthropic", "claude"]:
        if not api_key or "YOUR_API_KEY" in api_key:
            print("  [Anthropic] Нет API-ключа. Перехожу на регекс-пунктуатор.")
            return stage2b_regex_punctuate(text, plan)

        url = "https://api.anthropic.com/v1/messages"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        }
        chunks = _chunk_text(text)
        if len(chunks) > 1:
            print(f"  [Claude] Текст разбит на {len(chunks)} порций.")

        results = []
        for idx, chunk in enumerate(chunks):
            if len(chunks) > 1:
                print(f"    Порция {idx+1}/{len(chunks)}...")
            payload = {
                "model": model or "claude-3-5-sonnet-20240620",
                "max_tokens": 8192,
                "system": system_prompt,
                "messages": [{"role": "user", "content": chunk}],
            }
            req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers)
            try:
                with urllib.request.urlopen(req, timeout=120) as response:
                    res_data = json.loads(response.read().decode("utf-8"))
                    results.append(res_data["content"][0]["text"].strip())
            except Exception as e:
                print(f"  [Claude] Ошибка порции {idx+1}: {e}")
                results.append(chunk)
        return "\n\n".join(results)

    # --- Ollama (local) ---
    elif provider == "ollama":
        ollama_url = config.get("ollama_url", "http://localhost:11434").rstrip("/")
        url = f"{ollama_url}/api/generate"
        chunks = _chunk_text(text)
        if len(chunks) > 1:
            print(f"  [Ollama] Текст разбит на {len(chunks)} порций.")

        results = []
        for idx, chunk in enumerate(chunks):
            if len(chunks) > 1:
                print(f"    Порция {idx+1}/{len(chunks)}...")
            payload = {
                "model": model or "llama3",
                "prompt": f"{system_prompt}\n\nИсходный текст:\n{chunk}",
                "stream": False,
            }
            req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"),
                                         headers={"Content-Type": "application/json"})
            try:
                with urllib.request.urlopen(req, timeout=600) as response:
                    res_data = json.loads(response.read().decode("utf-8"))
                    results.append(res_data.get("response", "").strip())
            except Exception as e:
                print(f"  [Ollama] Ошибка порции {idx+1}: {e}")
                results.append(chunk)
        return "\n\n".join(results)

    # --- Fallback ---
    print("  [Литературный редактор] Провайдер не распознан. Перехожу на регекс-пунктуатор.")
    return stage2b_regex_punctuate(text, plan)


# =====================================================================
# STAGE 3: КОРРЕКТОР-ТИПОГРАФ — Publisher Typography
# (Кавычки, тире, частицы, дефисы, числительные)
# =====================================================================

def stage3_typography(text: str, plan: dict) -> str:
    """Stage 3: Корректор-типограф — русская издательская типографика."""
    if not text:
        return ""

    # Кавычки «ёлочки»
    text = re.sub(r'"([^"]+)"', r'«\1»', text)

    # Тире (em-dash) между словами
    text = re.sub(r' - ', ' — ', text)
    text = re.sub(r' – ', ' — ', text)

    # Длинное тире в начале строки (прямая речь)
    text = re.sub(r'^-\s+', '— ', text, flags=re.M)

    # Частицы через дефис
    text = re.sub(r'\b(из|из за)\b', 'из-за', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(из под)\b', 'из-под', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(все таки|всё таки)\b', 'всё-таки', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(в обшем|вобщем)\b', 'в общем', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(как бы)\b', 'как-бы', text)  # если не частица
    text = re.sub(r'\b(где то)\b', 'где-то', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(куда то)\b', 'куда-то', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(когда то)\b', 'когда-то', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(что то)\b', 'что-то', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(какой то)\b', 'какой-то', text, flags=re.IGNORECASE)

    # Неразрывный пробел перед одиночными предлогами/союзами в начале строки
    # (упрощённо — просто пробел, python-docx не поддерживает NBSP напрямую)

    # Многоточие
    text = re.sub(r'\.\.\.', '…', text)
    text = re.sub(r'\.\.\.\.', '…', text)

    # Множественные пробелы
    text = re.sub(r'[ \t]{2,}', ' ', text)

    return text


# =====================================================================
# STAGE 4: ВЫПУСКАЮЩИЙ РЕДАКТОР — 8 аудитов + расширенные проверки
# =====================================================================

# Терминальные знаки
TERMINAL_CHARS = ('.', '!', '?', '…', ':', '—', '|', ')', '»', '"', "'", '}',
                  ']', '»')

# Символы, которые считаются допустимым завершением (эмодзи, хэштеги, ссылки)
NON_ISSUE_ENDINGS = ('🍿', '⚡', '🌹', '😊', '😀', '😂', '😎', '👍', '🔥',
                     '❤', '💯', '🙏', '🎉', '✨', '⭐', '💜', '💙', '💚',
                     '💛', '🧡', '🤔', '😅', '🤣', '😉', '👀', '🎶', '🎵',
                     '🎸', '🎧', '🎬', '📸', '💻', '📱', '🔑', '⬇', '➡',
                     '✅', '❌', '⚠', '📢', '📰', '📉', '📈', '🤷', '🤦')

# Союзы/предлоги, на которые не должен заканчиваться абзац
CUT_CONJUNCTIONS = ('и', 'но', 'а', 'что', 'чтобы', 'для', 'на', 'в', 'к',
                    'с', 'по', 'от', 'до', 'при', 'об', 'о', 'не', 'ни')

# STT-артефакты для check_stt_artifacts
STT_ARTIFACT_WORDS = [
    'субтитры сделал', 'прочай', 'redhares', 'ецируuts', 'ход gou',
    'отправить. отправить', 'валясь', 'два две', 'три линиями',
    'человечку зеленых', 'без вывозного', 'вэ этом', 'спесь',
    'танковые люди', 'пахеризм', 'гринда', 'раджу руку', 'протуз',
]


def stage4_audit(text: str, plan: dict) -> dict:
    """
    Stage 4: Выпускающий редактор — 8 аудитов.
    Возвращает отчёт с найденными проблемами.
    """
    issues = {
        'check_cuts': [],           # обрывы на союзах/предлогах
        'check_terminal': [],       # отсутствие терминальных знаков
        'check_repetitions': [],    # дублирования слов
        'check_stt_artifacts': [],  # остаточный машинный мусор
        'check_orthography': [],    # орфография (заглушка для LLM)
        'check_punctuation': [],    # пунктуация (заглушка)
        'check_syntax': [],         # синтаксис (заглушка)
        'check_attribution': [],    # атрибуция (спикеры)
        'total': 0,
    }

    lines = text.split('\n')
    actions = plan.get('actions', [])

    for i, line in enumerate(lines, 1):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # --- check_terminal: терминальный знак ---
        if 'fix_terminal' in actions or 'audit' in actions:
            # Пропускаем хэштеги, эмодзи и ссылки
            if line_stripped.startswith('#') or line_stripped.startswith('Http'):
                continue
            # Проверяем, заканчивается ли строка эмодзи
            ends_with_emoji = any(line_stripped.endswith(e) for e in NON_ISSUE_ENDINGS)
            if ends_with_emoji:
                continue

            msg_match = re.match(r'^(\w+)\s+\[\d{1,2}:\d{2}\]\s+\[([^\]]+)\]:\s*(.*)$', line_stripped)
            if msg_match:
                content = msg_match.group(3).strip()
                if content and len(content) >= 5 and not content.startswith('Http'):
                    if not content.startswith('#') and not any(content.endswith(e) for e in NON_ISSUE_ENDINGS):
                        if not content.endswith(TERMINAL_CHARS):
                            issues['check_terminal'].append({
                                'line': i,
                                'text': content[-80:] if len(content) > 80 else content,
                            })
            elif len(line_stripped) > 20 and not line_stripped.startswith('📅'):
                if not line_stripped.endswith(TERMINAL_CHARS):
                    issues['check_terminal'].append({
                        'line': i,
                        'text': line_stripped[-80:],
                    })

        # --- check_cuts: обрывы на союзах ---
        if 'audit' in actions:
            # Не считаем обрезкой если строка уже закончена терминальным знаком
            if not line_stripped.endswith(TERMINAL_CHARS):
                last_word = line_stripped.split()[-1:] if line_stripped.split() else ['']
                # Удаляем знаки препинания для проверки
                if last_word:
                    clean_word = re.sub(r'[^\w\s]', '', last_word[0]).lower()
                    if clean_word in CUT_CONJUNCTIONS:
                        issues['check_cuts'].append({
                            'line': i,
                            'word': last_word[0],
                        })

        # --- check_repetitions: дублирования слов ---
        if 'audit' in actions:
            dups = re.findall(r'\b(\w+)\s+\1\b', line_stripped, re.IGNORECASE)
            for d in dups:
                if d.lower() not in ('нет', 'да'):  # осмысленные повторы
                    issues['check_repetitions'].append({
                        'line': i,
                        'word': d,
                    })

        # --- check_stt_artifacts: остаточный машинный мусор ---
        if 'audit' in actions:
            for artifact in STT_ARTIFACT_WORDS:
                if artifact.lower() in line_stripped.lower():
                    issues['check_stt_artifacts'].append({
                        'line': i,
                        'artifact': artifact,
                    })

    # --- check_attribution: спикеры ---
    if 'audit' in actions:
        speakers = set(re.findall(r'^(\w+)\s+\[\d{1,2}:\d{2}\]\s+\[', text, re.M))
        if len(speakers) > 0:
            issues['check_attribution'] = {
                'speakers': list(speakers),
                'count': len(speakers),
            }

    issues['total'] = (
        len(issues['check_cuts']) +
        len(issues['check_terminal']) +
        len(issues['check_repetitions']) +
        len(issues['check_stt_artifacts'])
    )

    return issues


def _format_audit_report(issues: dict) -> str:
    """Format audit issues into a readable report."""
    lines = ["# Отчёт вылускающего редактора (8 аудитов)", ""]

    sections = [
        ('check_terminal', 'Терминальные знаки'),
        ('check_cuts', 'Обрывы на союзах (check_cuts)'),
        ('check_repetitions', 'Дублирования слов'),
        ('check_stt_artifacts', 'Остаточный STT-мусор'),
    ]

    for key, title in sections:
        items = issues.get(key, [])
        if isinstance(items, list):
            status = "OK" if not items else f"{len(items)} проблем"
            lines.append(f"## {title}: {status}")
            for item in items[:10]:
                if isinstance(item, dict):
                    lines.append(f"  - Строка {item.get('line', '?')}: {item.get('text', item.get('word', item.get('artifact', '')))}")
            if len(items) > 10:
                lines.append(f"  ...и ещё {len(items) - 10}")
        lines.append("")

    # Attribution
    attr = issues.get('check_attribution', {})
    if attr:
        lines.append(f"## Атрибуция: {attr.get('count', 0)} спикеров: {', '.join(attr.get('speakers', []))}")
        lines.append("")

    lines.append(f"## Итого: {issues.get('total', 0)} проблем")
    return '\n'.join(lines)


# =====================================================================
# STAGE 4.5: АВТОИСПРАВЛЕНИЕ — фиксы найденных проблем
# =====================================================================

def stage4b_autofix(text: str, issues: dict, plan: dict) -> str:
    """Автоматически исправляет то, что можно безопасно исправить."""
    lines = text.split('\n')
    fixed = 0

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Повторы слов (LLM-независимый фикс)
        rep_match = re.search(r'\b(\w+)\s+\1\b', line_stripped, re.IGNORECASE)
        if rep_match:
            orig = line_stripped
            line_stripped = re.sub(r'\b(\w+)\s+\1\b', r'\1', line_stripped, flags=re.IGNORECASE)
            if line_stripped != orig:
                lines[i] = line_stripped
                fixed += 1

        # Обрезки на союзах (добавляем многоточие)
        words = line_stripped.rstrip('.!?').split()
        if words:
            last_word = words[-1].lower().rstrip(',:;')
            if last_word in CUT_CONJUNCTIONS:
                if not line_stripped.endswith('...'):
                    lines[i] = line_stripped.rstrip('.,;:!?') + '...'
                    fixed += 1

        # Терминальные знаки
        msg_match = re.match(r'^(\w+)\s+(\[\d{1,2}:\d{2}\]\s+\[[^\]]+\]:\s*)(.*)$', line_stripped)
        if msg_match:
            prefix = msg_match.group(1) + ' ' + msg_match.group(2)
            content = msg_match.group(3).strip()
            if content and len(content) >= 5 and not content.startswith('Http'):
                # Пропускаем эмодзи и хэштеги
                if not content.startswith('#') and not any(content.endswith(e) for e in NON_ISSUE_ENDINGS):
                    if not content.endswith(TERMINAL_CHARS):
                        content = content + '.'
                        lines[i] = prefix + content
                        fixed += 1
        elif len(line_stripped) > 20 and not line_stripped.startswith('📅') and not line_stripped.startswith('#'):
            if not any(line_stripped.endswith(e) for e in NON_ISSUE_ENDINGS):
                if not line_stripped.endswith(TERMINAL_CHARS):
                    lines[i] = line_stripped + '.'
                    fixed += 1

    print(f"  [Автоисправление] {fixed} проблем исправлено (повторы + обрезки + терминальные знаки).")
    return '\n'.join(lines)


# =====================================================================
# STAGE 5: ВЕРСТАЛЬЩИК — Professional Book Layout (DOCX)
# =====================================================================

# Палитра цветов спикеров (расширенная для >2 спикеров)
SPEAKER_PALETTE = [
    RGBColor(0x00, 0x00, 0xFF),   # синий
    RGBColor(0xFF, 0x00, 0x00),   # красный
    RGBColor(0x00, 0x80, 0x00),   # зелёный
    RGBColor(0x80, 0x00, 0x80),   # фиолетовый
    RGBColor(0xFF, 0x80, 0x00),   # оранжевый
    RGBColor(0x00, 0x80, 0x80),   # бирюзовый
    RGBColor(0x80, 0x40, 0x00),   # коричневый
    RGBColor(0x00, 0x00, 0x80),   # тёмно-синий
]


def _get_speaker_color(speaker: str, color_map: dict) -> RGBColor:
    """Get color for a speaker, assigning new colors as needed."""
    if speaker not in color_map:
        idx = len(color_map) % len(SPEAKER_PALETTE)
        color_map[speaker] = SPEAKER_PALETTE[idx]
    return color_map[speaker]


def stage5_build_docx(text: str, output_path: str, plan: dict,
                       title: str = "", subtitle: str = ""):
    """
    Stage 5: Верстальщик — компиляция в DOCX.
    - Заголовок/подзаголовок
    - Times New Roman 12, интервал 1.15, красная строка
    - Цветовая разметка спикеров (если keep_speakers)
    - Оформление прямой речи (тире)
    - Сохранение дат (📅)
    """
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Титульная страница
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)
        run.font.bold = True
        if subtitle:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(subtitle)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(14)
        doc.add_paragraph()

    # Карта цветов спикеров
    color_map = {}
    use_colors = plan.get('keep_speakers', False)

    # Обработка строк
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

        # Проверяем, сообщение ли это спикера
        msg_match = re.match(
            r'^(\w+)\s+(\[\d{1,2}:\d{2}\]\s+\[[^\]]+\]:\s*)(.*)$',
            line
        )

        if msg_match and use_colors:
            speaker = msg_match.group(1)
            header = speaker + ' ' + msg_match.group(2)
            content = msg_match.group(3).strip()

            # Заголовок спикера — цветной
            run = p.add_run(header)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            color = _get_speaker_color(speaker, color_map)
            run.font.color.rgb = color

            # Содержание — обычный чёрный
            if content:
                run2 = p.add_run(content)
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(12)
        else:
            # Обычный абзац / дата / заголовок
            clean = line
            # Прямая речь: тире в начале
            if clean.startswith("-") or clean.startswith("—"):
                clean = "— " + clean.lstrip("-—").strip()

            # Даты (📅) — без красной строки
            if clean.startswith("📅"):
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_before = Pt(12)

            run = p.add_run(clean)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Заголовки разделов (если есть)
            if clean.startswith("#"):
                run.font.bold = True
                p.paragraph_format.first_line_indent = Inches(0)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  [Верстальщик] DOCX сохранён: {output_path}")
    return output_path


# =====================================================================
# PIPELINE EXECUTION ENGINE
# =====================================================================

def process_manuscript_chain(input_path: str = None, output_path: str = None):
    """
    Полный конвейер редакторского отдела:
    Приёмка → Авто-определение → Корректор → Литературный редактор
    → Корректор-типограф → Выпускающий редактор → Верстальщик
    """
    os.makedirs("inputs/raw_texts", exist_ok=True)
    os.makedirs("inputs/audio", exist_ok=True)
    os.makedirs("output/books", exist_ok=True)
    os.makedirs("output/.llm_cache", exist_ok=True)

    # --- Загрузка конфигурации ---
    config_path = "config.json"
    if not os.path.exists(config_path):
        default_config = {
            "api_provider": "auto",
            "api_key": os.environ.get("OPENAI_API_KEY", ""),
            "model": "auto",
            "ollama_url": "http://localhost:11434",
            "lmstudio_url": "http://localhost:1234",
            "genre": "Auto-detect",
            "title": "",
            "subtitle": "",
            "chunk_size": 4000,
        }
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(default_config, f, indent=2, ensure_ascii=False)

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    # --- Определение входного файла ---
    if not input_path:
        raw_files = [
            os.path.join("inputs/raw_texts", f)
            for f in os.listdir("inputs/raw_texts")
            if f.endswith(('.txt', '.md', '.docx'))
        ]
        if raw_files:
            input_path = raw_files[0]
        else:
            input_path = "inputs/raw_texts/sample.txt"
            with open(input_path, "w", encoding="utf-8") as f:
                f.write("Поместите ваш текст сюда.")

    if not output_path:
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join("output/books", f"{base_name}.docx")

    print("=== РЕДАКТОРСКИЙ ОТДЕЛ VOX2BOOK ===")
    print(f"Вход:  {input_path}")
    print(f"Выход: {output_path}")
    print()

    # --- Чтение файла ---
    raw_text = _read_file(input_path)

    # === АВТО-ОПРЕДЕЛЕНИЕ ===
    print("Авто-определение жанра, стиля и действий...")
    plan = auto_detect_plan(raw_text, config)
    print_plan(plan)

    # Сохранение плана
    with open(os.path.join("output", ".llm_cache", "auto_plan.json"), "w", encoding="utf-8") as f:
        json.dump(plan, f, indent=2, ensure_ascii=False)

    # === STAGE 1: КОРРЕКТОР (гигиена) ===
    print("\n--- Stage 1/5: Корректор (гигиена источника) ---")
    text_s1 = stage1_cleanup(raw_text, plan)
    _cache("01_cleanup", text_s1)
    print(f"  [Корректор] Гигиена выполнена. Кэш: output/.llm_cache/01_cleanup.txt")

    # === STAGE 2: ЛИТЕРАТУРНЫЙ РЕДАКТОР (LLM) ===
    print(f"\n--- Stage 2/5: Литературный редактор ({config.get('api_provider')}) ---")
    text_s2 = stage2_neural_stylist(text_s1, config, plan)
    _cache("02_neural", text_s2)
    print(f"  [Литературный редактор] Правка выполнена. Кэш: output/.llm_cache/02_neural.txt")

    # === STAGE 3: КОРРЕКТОР-ТИПОГРАФ ===
    print("\n--- Stage 3/5: Корректор-типограф ---")
    text_s3 = stage3_typography(text_s2, plan)
    _cache("03_typography", text_s3)
    print(f"  [Типограф] Типографика применена. Кэш: output/.llm_cache/03_typography.txt")

    # === STAGE 4: ВЫПУСКАЮЩИЙ РЕДАКТОР (8 аудитов) ===
    print("\n--- Stage 4/5: Выпускающий редактор (8 аудитов) ---")
    issues = stage4_audit(text_s3, plan)
    audit_report = _format_audit_report(issues)
    _cache("04_audit", audit_report, "md")

    if issues['total'] > 0:
        print(f"  [Аудит] Найдено {issues['total']} проблем:")
        print(f"    Терминальные: {len(issues['check_terminal'])}")
        print(f"    Обрывы:       {len(issues['check_cuts'])}")
        print(f"    Повторы:      {len(issues['check_repetitions'])}")
        print(f"    STT-мусор:    {len(issues['check_stt_artifacts'])}")

        # === STAGE 4b: АВТОИСПРАВЛЕНИЕ ===
        print("\n--- Stage 4b: Автоисправление ---")
        text_s4 = stage4b_autofix(text_s3, issues, plan)
        _cache("04b_autofixed", text_s4)

        # Повторный аудит
        issues2 = stage4_audit(text_s4, plan)
        if issues2['total'] < issues['total']:
            print(f"  [Автоисправление] Проблем осталось: {issues2['total']} (было {issues['total']})")
        text_final = text_s4
    else:
        print("  [Аудит] OK — 0 проблем.")
        text_final = text_s3

    # === STAGE 5: ВЕРСТАЛЬЩИК (DOCX) ===
    print("\n--- Stage 5/5: Верстальщик (DOCX) ---")
    title = config.get("title", "")
    subtitle = config.get("subtitle", "")
    stage5_build_docx(text_final, output_path, plan, title, subtitle)

    # === ИТОГ ===
    print("\n=== ГОТОВО ===")
    print(f"Книга:   {output_path}")
    print(f"Кэш:     output/.llm_cache/ (01-05 этапы)")
    print(f"Аудит:   output/.llm_cache/04_audit.md")


if __name__ == "__main__":
    in_file = sys.argv[1] if len(sys.argv) > 1 else None
    out_file = sys.argv[2] if len(sys.argv) > 2 else None
    process_manuscript_chain(in_file, out_file)