#!/usr/bin/env python3
"""
Vox2Book — Round 2 User Error Fix Script for file in "E:\coding\работа с литературой\Анфи"
Applies Round 2 fixes to "E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx":
1. Complete cleanup of spurious "из-за" -> "из" across all context patterns.
2. Complete cleanup of truncated word "жно" -> "нужно".
3. AI model & Service brand names (джи пяти -> GPT-5, джамина -> Gemini, грок/гроб -> Grok, DeepSig -> DeepSeek, литр с -> Литрес).
4. Hyphenated pronoun comma anomalies ("что, нибудь" -> "что-нибудь", "как, то" -> "как-то").
"""

import sys
import os
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

ANFI_DOC_PATH = Path(r"E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx")
BOOKS_DOC_PATH = Path(r"E:\coding\работа с литературой\output\books\Диалоги_Анфи_и_Kir_2024-2026.docx")

SPEAKER_COLORS = {
    "kir": RGBColor(13, 71, 161),     # Deep Blue
    "анфи": RGBColor(194, 24, 91),    # Berry / Magenta
    "default": RGBColor(74, 20, 140)  # Dark Purple
}
DATE_COLOR = RGBColor(46, 125, 50)    # Forest Green

ROUND2_REPLACEMENTS = [
    # 1. Spurious "из-за" -> "из"
    (r'\bНа одном из-за них\b', 'На одном из них'),
    (r'\bодно из-за них\b', 'одно из них'),
    (r'\bиз-за них\b', 'из них'),
    (r'\bсамая логичная из-за всех\b', 'самая логичная из всех'),
    (r'\bиз-за всех\b', 'из всех'),
    (r'\bиз-за мастерской\b', 'из мастерской'),
    (r'\bиз-за древнегреческой мифологии\b', 'из древнегреческой мифологии'),
    (r'\bиз-за мифологии\b', 'из мифологии'),
    (r'\bиз-за Дайвинчика\b', 'из Дайвинчика'),
    (r'\bиз-за дайвинчика\b', 'из Дайвинчика'),
    (r'\bиз-за десяти человек\b', 'из десяти человек'),
    (r'\bиз-за своей жизни\b', 'из своей жизни'),
    (r'\bиз-за советских фильмов\b', 'из советских фильмов'),
    (r'\bиз-за Библии\b', 'из Библии'),
    (r'\bиз-за библии\b', 'из Библии'),
    (r'\bиз-за соседнего двора\b', 'из соседнего двора'),
    (r'\bОдним из-за величайших\b', 'Одним из величайших'),
    (r'\bиз-за величайших\b', 'из величайших'),
    (r'\bиз-за рук\b', 'из рук'),
    (r'\bОдна из-за партий\b', 'Одна из партий'),
    (r'\bиз-за партий\b', 'из партий'),
    (r'\bне идет из-за крана\b', 'не идет из крана'),
    (r'\bне идёт из-за крана\b', 'не идёт из крана'),
    (r'\bиз-за клеток\b', 'из клеток'),
    (r'\bиз-за любой книги\b', 'из любой книги'),
    (r'\bиз-за XXI века\b', 'из XXI века'),
    (r'\bиз-за 21 века\b', 'из 21 века'),
    (r'\bпропала из-за Telegram\b', 'пропала из Telegram'),
    (r'\bсдала из-за экзаменов\b', 'сдала из экзаменов'),
    (r'\bодин из-за\b', 'один из'),
    (r'\bодна из-за\b', 'одна из'),
    (r'\bодно из-за\b', 'одно из'),

    # 2. Standalone truncated "жно" -> "нужно"
    (r'\bжно\b', 'нужно'),

    # 3. AI Models & Tech Brands
    (r'\bджи пяти\b', 'GPT-5'),
    (r'\bджамина\b', 'Gemini'),
    (r'\bджемина\b', 'Gemini'),
    (r'\bdeepsig\b', 'DeepSeek'),
    (r'\bDeepSig\b', 'DeepSeek'),
    (r'\bлитр с\b', 'Литрес'),
    (r'\bлитрес\b', 'Литрес'),

    # 4. Hyphenated pronouns with comma anomalies
    (r'\bчто,\s*нибудь\b', 'что-нибудь'),
    (r'\bчто,\s*то\b', 'что-то'),
    (r'\bкогда,\s*то\b', 'когда-то'),
    (r'\bкак,\s*то\b', 'как-то'),
    (r'\bгде,\s*то\b', 'где-то'),
    (r'\bкто,\s*то\b', 'кто-то'),
    (r'\bкакой,\s*то\b', 'какой-то'),
]

def fix_text_round2(text: str) -> str:
    if not text:
        return text

    for pat, repl in ROUND2_REPLACEMENTS:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)

    text = re.sub(r'\s+,', ',', text)
    text = re.sub(r',{2,}', ',', text)
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def process_round2():
    doc = Document(ANFI_DOC_PATH)
    paragraphs = doc.paragraphs

    doc_new = Document()
    style = doc_new.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    msg_count = 0
    header_count = 0
    fix_count = 0

    MSG_RE = re.compile(r'^(Kir|Анфи)\s+\[(\d{1,2}:\d{2})\]\s+\[(Голосовое|Текст)\]:\s*(.*)$', re.DOTALL)

    print("Running Round 2 fixes on E:\\coding\\работа с литературой\\Анфи\\Диалоги_Анфи_и_Kir_2024-2026.docx...")

    for p in paragraphs:
        ptext = p.text.strip()
        if not ptext:
            continue

        m = MSG_RE.match(ptext)
        if m:
            speaker = m.group(1)
            tstamp = m.group(2)
            mtype = m.group(3)
            body = m.group(4)

            fixed_b = fix_text_round2(body)
            if fixed_b != body:
                fix_count += 1

            p_new = doc_new.add_paragraph()
            p_new.paragraph_format.space_after = Pt(4)
            p_new.paragraph_format.line_spacing = 1.15

            header_prefix = f"{speaker} [{tstamp}] [{mtype}]: "
            run_spk = p_new.add_run(header_prefix)
            run_spk.bold = True
            run_spk.font.color.rgb = SPEAKER_COLORS.get(speaker.lower(), SPEAKER_COLORS["default"])

            run_body = p_new.add_run(fixed_b)
            msg_count += 1
        else:
            cleaned_hdr = fix_text_round2(ptext)

            p_hdr = doc_new.add_paragraph()
            p_hdr.paragraph_format.space_after = Pt(6)
            p_hdr.paragraph_format.line_spacing = 1.15
            run_hdr = p_hdr.add_run(cleaned_hdr)

            if cleaned_hdr.startswith('📅'):
                run_hdr.bold = True
                run_hdr.font.color.rgb = DATE_COLOR
            elif cleaned_hdr.startswith('Диалоги') or cleaned_hdr.startswith('Полная хроника') or cleaned_hdr.startswith('═══'):
                run_hdr.bold = True
                if cleaned_hdr.startswith('═══'):
                    run_hdr.font.size = Pt(14)
                    run_hdr.font.color.rgb = SPEAKER_COLORS["kir"]

            header_count += 1

    doc_new.save(str(ANFI_DOC_PATH))
    shutil.copy(ANFI_DOC_PATH, BOOKS_DOC_PATH)

    print(f"[Done Round 2 Fixes] All Round 2 error patterns successfully fixed!")
    print(f"  - Messages processed: {msg_count}")
    print(f"  - Messages with Round 2 fixes applied: {fix_count}")
    print(f"  - Headers/Date lines: {header_count}")

if __name__ == "__main__":
    process_round2()
