import docx
import re

def fix_stt_errors(text):
    fixes = {
        r'расставлю температура': 'валяюсь, температура',
        r'недосыб': 'недосып',
        r'Голосовое:\s*Hello': 'Голосовое: Привет',
    }
    for pattern, replacement in fixes.items():
        text = re.sub(pattern, replacement, text)
    return text

def fix_sentence_endings(text):
    text = re.sub(r'([А-ЯЁа-яё])([А-ЯЁа-яё])\s+([А-ЯЁа-яё])', r'\1\2. \3', text)
    text = re.sub(r'([.!?…])\s+([А-ЯЁа-яё])([А-ЯЁа-яё])\s+([А-ЯЁа-яё])', r'\1 \2\3. \4', text)
    return text

def process_dialogue(input_path, output_path):
    doc = docx.Document(input_path)
    
    new_doc = docx.Document()
    new_doc.add_paragraph('Диалоги и устная речь')
    new_doc.add_paragraph('')
    new_doc.add_paragraph('Полная хроника общения (2026 г.)')
    new_doc.add_paragraph('Собеседники: Kir и Анфи')
    
    current_date = None
    
    for para in doc.paragraphs:
        text = para.text
        
        date_match = re.match(r'📅\s+(.+?)\s+г\.', text)
        if date_match:
            current_date = date_match.group(1)
            new_doc.add_paragraph(f'📅 {current_date} г.')
            continue
        
        if current_date is None:
            continue
        
        if '📅' in text and not date_match:
            continue
        
        if re.match(r'^(Анфи|Kir)\s+\[', text):
            text = fix_stt_errors(text)
            text = fix_sentence_endings(text)
            new_doc.add_paragraph(text)
        elif text and not text.startswith('📅'):
            if 'https://' not in text.lower() and 'http://' not in text.lower():
                if not text.startswith('👌') and not text.startswith('🍿'):
                    if not text.startswith('H') or 'ttps' not in text:
                        if not text.startswith('Z') or 'RRF' not in text:
                            if not text.startswith('K') or '.com' not in text:
                                if not text.startswith('A') or 'ttps' not in text:
                                    if not text.startswith('П') or 'риветствуйте' not in text:
                                        new_doc.add_paragraph(text)
    
    new_doc.save(output_path)
    print(f'Processed {len(new_doc.paragraphs)} paragraphs')

if __name__ == "__main__":
    input_file = r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx'
    output_file = r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx'
    process_dialogue(input_file, output_file)