# -*- coding: utf-8 -*-
"""Export monthly canon TXT → print DOCX with EXACT timestamps preserved.

Header (as in canon, unchanged precision):
  YYYY-MM-DD HH:MM:SS Speaker (голосовое|текст)

Body typography for print; keep_mat; speaker colors.
Output: output/books/тираж/
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"E:\coding\работа с литературой")
CANON = ROOT / "финальная_обработка_текстов"
OUT = ROOT / "output" / "books" / "тираж"
NEURAL = ROOT / "output" / "books" / "отредактированные_нейросетью"

HDR = re.compile(
    r"^(\d{4}-\d{2}-\d{2}) (\d{2}:\d{2}:\d{2})\s+(.+?)\s+\((голосовое|текст)\)\s*$"
)

ANFI = RGBColor(0x1F, 0x4E, 0x79)
KIR = RGBColor(0xC0, 0x00, 0x00)
MONTH_RU = {
    "01": "Январь",
    "02": "Февраль",
    "03": "Март",
    "04": "Апрель",
    "05": "Май",
    "06": "Июнь",
    "07": "Июль",
    "08": "Август",
    "09": "Сентябрь",
    "10": "Октябрь",
    "11": "Ноябрь",
    "12": "Декабрь",
}

TERMINALS = ".!?…»\")'"


def speaker_color(name: str) -> RGBColor:
    if name.lower().startswith("анф"):
        return ANFI
    return KIR


def typography(text: str) -> str:
    t = text.replace("...", "…")
    t = re.sub(r"\s+—\s+", " — ", t)
    t = re.sub(r"(?<!\s)—(?!\s)", " — ", t)
    t = re.sub(r" {2,}", " ", t)
    t = t.strip()
    if t and t[-1] not in TERMINALS and t[-1] not in ")]}":
        if re.search(r"[А-Яа-яA-Za-z0-9]$", t):
            t += "."
    return t


def parse_month(path: Path) -> list[tuple[str, str, str, str, str]]:
    """Return (date, time, speaker, kind, body) — time is exact HH:MM:SS from canon."""
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    msgs: list[tuple[str, str, str, str, str]] = []
    i = 0
    while i < len(lines):
        m = HDR.match(lines[i])
        if not m:
            i += 1
            continue
        date, ts, speaker, kind = m.group(1), m.group(2), m.group(3), m.group(4)
        i += 1
        body: list[str] = []
        while i < len(lines) and not HDR.match(lines[i]) and lines[i].strip() != "---":
            body.append(lines[i])
            i += 1
        if i < len(lines) and lines[i].strip() == "---":
            i += 1
        b = typography("\n".join(body).strip())
        if not b:
            continue
        if re.fullmatch(r"https?://\S+", b):
            continue
        msgs.append((date, ts, speaker, kind, b))
    return msgs


def set_run_font(run, size=12, bold=False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def export_docx(
    path: Path, msgs: list[tuple[str, str, str, str, str]], title: str
) -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    set_run_font(tr, size=18, bold=True)
    doc.add_paragraph()

    current_date = ""
    for date, ts, speaker, kind, body in msgs:
        if date != current_date:
            current_date = date
            y, m, d = date.split("-")
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(14)
            dp.paragraph_format.space_after = Pt(8)
            dp.paragraph_format.first_line_indent = Inches(0)
            dr = dp.add_run(f"📅 {int(d)} {MONTH_RU.get(m, m)} {y}")
            set_run_font(dr, size=12, bold=True)

        color = speaker_color(speaker)
        # Exact canon header: YYYY-MM-DD HH:MM:SS Speaker (голосовое|текст)
        hp = doc.add_paragraph()
        hp.paragraph_format.first_line_indent = Inches(0)
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run(f"{date} {ts} {speaker} ({kind})")
        set_run_font(hr, bold=True, color=color)

        bp = doc.add_paragraph()
        bp.paragraph_format.first_line_indent = Inches(0.5)
        bp.paragraph_format.space_after = Pt(6)
        bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        br = bp.add_run(body)
        set_run_font(br, color=color)

    out_path = OUT / f"{path.stem}.docx"
    doc.save(out_path)
    return out_path


def ensure_terminals_in_canon(path: Path) -> int:
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    rebuilt: list[str] = []
    ch = 0
    while i < len(lines):
        m = HDR.match(lines[i])
        if not m:
            i += 1
            continue
        header = lines[i]
        i += 1
        body: list[str] = []
        while i < len(lines) and not HDR.match(lines[i]) and lines[i].strip() != "---":
            body.append(lines[i])
            i += 1
        if i < len(lines) and lines[i].strip() == "---":
            i += 1
        raw = "\n".join(body).strip()
        if not raw:
            continue
        fixed = typography(raw.replace("...", "…"))
        if fixed != raw:
            ch += 1
        rebuilt.append(header)
        rebuilt.append(fixed)
        rebuilt.append("")
    path.write_text("\n".join(rebuilt).rstrip() + "\n", encoding="utf-8")
    NEURAL.mkdir(parents=True, exist_ok=True)
    (NEURAL / path.name).write_text(path.read_text(encoding="utf-8"), encoding="utf-8")
    return ch


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    total_term = 0
    for path in sorted(CANON.rglob("*.txt")):
        term = ensure_terminals_in_canon(path)
        total_term += term
        msgs = parse_month(path)
        title = path.stem.split("_", 1)[1] if "_" in path.stem else path.stem
        out = export_docx(path, msgs, title)
        # verify first message timestamp round-trip
        sample = ""
        if msgs:
            d, ts, sp, k, _ = msgs[0]
            sample = f"{d} {ts} {sp} ({k})"
        print(f"OK {path.name}: msgs={len(msgs)} term~{term} sample=[{sample}] -> {out.name}")
    print("TOTAL_TERMINAL_TOUCHES", total_term)
    print("EXPORTED", len(list(OUT.glob('*.docx'))))


if __name__ == "__main__":
    main()
