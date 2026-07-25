# -*- coding: utf-8 -*-
import docx
import re
from docx.shared import RGBColor

doc = docx.Document(r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx')

ANFI_COLOR = RGBColor(0x00, 0x00, 0xFF)
KIR_COLOR = RGBColor(0xFF, 0x00, 0x00)

fixes = [
    ('Поздравляю с 8 Марта! 🌹', 'Поздравляю с 8 Марта! 🌹.'),
    ('Ставишь вот', 'Ставишь вот.'),
    ('– Активирован Prime', '– Активирован Prime.'),
]

fixed = 0
for para in doc.paragraphs:
    text = para.text
    msg_match = re.match(r'^(Анфи|Kir)\s+\[(\d{2}:\d{2})\]\s+\[([^\]]+)\]:\s*(.*)$', text)
    if not msg_match:
        continue
    
    speaker, time, msg_type, content = msg_match.groups()
    changed = False
    
    for old_text, new_text in fixes:
        if old_text in content:
            content = content.replace(old_text, new_text)
            changed = True
    
    if changed:
        for run in para.runs:
            run.text = ''
        header = speaker + ' [' + time + '] [' + msg_type + ']: '
        color = ANFI_COLOR if speaker == 'Анфи' else KIR_COLOR
        if para.runs:
            run = para.runs[0]
            run.text = header
            run.font.color.rgb = color
        else:
            run = para.add_run(header)
            run.font.color.rgb = color
        para.add_run(content)
        fixed += 1

doc.save(r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx')
print('Исправлено:', fixed)