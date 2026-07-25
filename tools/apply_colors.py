import docx
from docx.shared import RGBColor
import re
import os

template_file = r'E:\coding\работа с литературой\output\books\Диалоги_Анфи_и_Kir_2024-2025.docx'
input_file = r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx'

input_doc = docx.Document(input_file)

# Colors: Анфи - синий (0x0000FF), Kir - красный (0xFF0000)
ANFI_COLOR = RGBColor(0x00, 0x00, 0xFF)  # Blue
KIR_COLOR = RGBColor(0xFF, 0x00, 0x00)   # Red

def process_paragraph(paragraph):
    text = paragraph.text
    
    # Match speaker patterns
    pattern = r'^(Анфи|Kir)\s+\[(\d{2}:\d{2})\]\s+\[([^\]]+)\]:\s*(.*)$'
    match = re.match(pattern, text)
    
    if match:
        speaker, time, msg_type, content = match.groups()
        
        # Clear runs and recreate with colors
        paragraph.clear()
        
        # Add speaker with color
        run = paragraph.add_run(f"{speaker} [{time}] [{msg_type}]: ")
        run.font.color.rgb = ANFI_COLOR if speaker == 'Анфи' else KIR_COLOR
        
        # Add content without color (or with default)
        run = paragraph.add_run(content)
        
    return paragraph

# Process all paragraphs
for i, para in enumerate(input_doc.paragraphs):
    if '[Голосовое]' in para.text or '[Текст]' in para.text:
        process_paragraph(para)

# Save
input_doc.save(input_file)
print(f"Processed {len(input_doc.paragraphs)} paragraphs with colored speakers")
print("Saved to:", input_file)