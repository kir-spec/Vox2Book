import docx
import re

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text

def process_message(line):
    match = re.match(r'^(Анфи|Kir)\s+\[(\d{2}:\d{2})\]\s+\[([^\]]+)\]:\s*(.*)$', line)
    if match:
        name, time, msg_type, content = match.groups()
        if name == "Анфи":
            name = "Анфи"
        else:
            name = "Kir"
        return f"{name} [{time}] [{msg_type}]: {clean_text(content)}"
    return None

def process_section(input_path, output_path, section_date):
    doc = docx.Document(output_path)
    doc.add_paragraph("")
    doc.add_paragraph(f"📅 {section_date}")
    
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_section = False
    for line in lines:
        line = line.strip()
        if section_date in line and '📅' in line:
            in_section = True
            continue
        if in_section:
            if line.startswith('📅') and section_date not in line:
                break
            if line and not line.startswith('📅'):
                processed = process_message(line)
                if processed:
                    doc.add_paragraph(processed)
    
    doc.save(output_path)
    print(f"Section {section_date} added to {output_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        section_date = sys.argv[3] if len(sys.argv) > 3 else "26 января 2026 г."
        process_section(input_file, output_file, section_date)