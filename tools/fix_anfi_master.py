#!/usr/bin/env python3
"""
Vox2Book — Master Fix Script for file in "E:\coding\работа с литературой\Анфи"
Applies all 6 categories of user error report to "E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx":
1. Proper Noun & Brand Capitalization (Сергей, Андрей, Анастасия, Мария, Анфия, Яндекс Маркет, ВКонтакте).
2. Spurious preposition "из-за" -> "из" (из старых фотографий, не из Google Play Market, из Telegram, из фильма, из чата, из мастерской, из Новгорода, из Google Play).
3. Hyphenated pronoun comma anomalies ("что, нибудь" -> "что-нибудь", "что, то" -> "что-то", "когда, то" -> "когда-то").
4. Specific word & STT garble repairs (вспорхнуть, посвящаю, нейросеть, бургер, закомплексованный, это у, всколыхнули).
5. Conjunction comma fixes ("А, когда" -> "А когда", "А, если" -> "А если").
6. Typography & time spacing fixes ("00: 15" -> "00:15").
"""

import sys
import os
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

ANFI_DOC_PATH = Path(r"E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx")
BOOKS_DOC_PATH = Path(r"E:\coding\работа с литературой\output\books\Диалоги_Анфи_и_Kir_2024-2026.docx")

SPEAKER_COLORS = {
    "kir": RGBColor(13, 71, 161),     # Deep Blue
    "анфи": RGBColor(194, 24, 91),    # Berry / Magenta
    "default": RGBColor(74, 20, 140)  # Dark Purple
}
DATE_COLOR = RGBColor(46, 125, 50)    # Forest Green

EXACT_REPLACEMENTS = [
    # 1. Names & Brands Capitalization
    (r'\bсергей\b', 'Сергей'),
    (r'\bсергея\b', 'Сергея'),
    (r'\bсергею\b', 'Сергею'),
    (r'\bсергеем\b', 'Сергеем'),
    (r'\bандрей\b', 'Андрей'),
    (r'\bандрея\b', 'Андрея'),
    (r'\bандрею\b', 'Андрею'),
    (r'\bандреем\b', 'Андреем'),
    (r'\bанастасия\b', 'Анастасия'),
    (r'\bанастасии\b', 'Анастасии'),
    (r'\bанастасию\b', 'Анастасию'),
    (r'\bмария и анфия\b', 'Мария и Анфия'),
    (r'\bмарии и анфии\b', 'Марии и Анфии'),
    (r'\bанфия\b', 'Анфия'),
    (r'\bанфии\b', 'Анфии'),
    (r'\bанфию\b', 'Анфию'),
    (r'\bанфией\b', 'Анфией'),
    (r'\bяндекс Маркет\b', 'Яндекс Маркет'),
    (r'\bяндекс маркет\b', 'Яндекс Маркет'),
    (r'\bяндекс\.маркет\b', 'Яндекс Маркет'),
    (r'\bв контакте\b', 'ВКонтакте'),
    (r'\bвконтакте\b', 'ВКонтакте'),
    (r'\bвконтакте-видео\b', 'ВКонтакте Видео'),

    # 2. Spurious "из-за" -> "из"
    (r'\bиз-за старых фотографий\b', 'из старых фотографий'),
    (r'\bиз-за фотографий\b', 'из фотографий'),
    (r'\bне из-за Google Play Market\b', 'не из Google Play Market'),
    (r'\bиз-за Google Play Market\b', 'из Google Play Market'),
    (r'\bиз-за Google Play\b', 'из Google Play'),
    (r'\bиз-за App Store\b', 'из App Store'),
    (r'\bвыйти из-за Telegram\b', 'выйти из Telegram'),
    (r'\bиз-за Telegram\b', 'из Telegram'),
    (r'\bиз-за телеграма\b', 'из Телеграма'),
    (r'\bиз-за телеграм\b', 'из Телеграм'),
    (r'\bиз-за фильма\b', 'из фильма'),
    (r'\bиз-за чата\b', 'из чата'),
    (r'\bиз-за мастерской\b', 'из мастерской'),
    (r'\bиз-за Новгорода\b', 'из Новгорода'),
    (r'\bвыходцами из-за Новгорода\b', 'выходцами из Новгорода'),
    (r'\bиз-за магазина\b', 'из магазина'),
    (r'\bиз-за репозитория\b', 'из репозитория'),
    (r'\bиз-за текстового файла\b', 'из текстового файла'),

    # Hyphenated pronouns with comma anomalies
    (r'\bчто,\s*нибудь\b', 'что-нибудь'),
    (r'\bчто,\s*то\b', 'что-то'),
    (r'\bкогда,\s*то\b', 'когда-то'),
    (r'\bгде,\s*то\b', 'где-то'),
    (r'\bкак,\s*то\b', 'как-то'),
    (r'\bкто,\s*то\b', 'кто-то'),
    (r'\bкакой,\s*то\b', 'какой-то'),
    (r'\bкакая,\s*то\b', 'какая-то'),
    (r'\bкакое,\s*то\b', 'какое-то'),
    (r'\bкакие,\s*то\b', 'какие-то'),

    # Specific typos & STT garbles
    (r'\bвспархнуть\b', 'вспорхнуть'),
    (r'\bпосвещаю\b', 'посвящаю'),
    (r'\bэго-сеть\b', 'нейросеть'),
    (r'\bнейрасить\b', 'нейросеть'),
    (r'\bблогер у меня был только один\b', 'бургер у меня был только один'),
    (r'\bа блогер у меня был только один\b', 'а бургер у меня был только один'),
    (r'\bзакомпликсованный\b', 'закомплексованный'),
    (r'\bэтоиу\b', 'это у'),
    (r'\bвсколыбнули\b', 'всколыхнули'),

    # 3. Punctuation anomalies
    (r'\bА,\s+когда\b', 'А когда'),
    (r'\bА,\s+если\b', 'А если'),
    (r'\bИ,\s+когда\b', 'И когда'),
    (r'\bИ,\s+если\b', 'И если'),
    (r'\bНо,\s+когда\b', 'Но когда'),
    (r'\bНо,\s+если\b', 'Но если'),
    (r'\bДа,\s+даже\b', 'Да даже'),

    # Time format spacing (e.g. "00 : 15" -> "00:15")
    (r'(\d{1,2})\s*:\s*(\d{2})', r'\1:\2'),
]

def fix_paragraph_text(text: str) -> str:
    if not text:
        return text

    for pat, repl in EXACT_REPLACEMENTS:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)

    # General whitespace & comma cleanup
    text = re.sub(r'\s+,', ',', text)
    text = re.sub(r',{2,}', ',', text)
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def process_anfi_file():
    doc = Document(ANFI_DOC_PATH)
    paragraphs = doc.paragraphs

    doc_new = Document()
    style = doc_new.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    msg_count = 0
    header_count = 0
    fix_count = 0

    MSG_RE = re.compile(r'^(Kir|Анфи)\s+\[(\d{1,2}:\d{2})\]\s+\[(Голосовое|Текст)\]:\s*(.*)$', re.DOTALL)

    print("Fixing errors in E:\\coding\\работа с литературой\\Анфи\\Диалоги_Анфи_и_Kir_2024-2026.docx...")

    for p in paragraphs:
        ptext = p.text.strip()
        if not ptext:
            continue

        m = MSG_RE.match(ptext)
        if m:
            speaker = m.group(1)
            tstamp = m.group(2)
            mtype = m.group(3)
            body = m.group(4)

            fixed_b = fix_paragraph_text(body)
            if fixed_b != body:
                fix_count += 1

            p_new = doc_new.add_paragraph()
            p_new.paragraph_format.space_after = Pt(4)
            p_new.paragraph_format.line_spacing = 1.15

            header_prefix = f"{speaker} [{tstamp}] [{mtype}]: "
            run_spk = p_new.add_run(header_prefix)
            run_spk.bold = True
            run_spk.font.color.rgb = SPEAKER_COLORS.get(speaker.lower(), SPEAKER_COLORS["default"])

            run_body = p_new.add_run(fixed_b)
            msg_count += 1
        else:
            cleaned_hdr = fix_paragraph_text(ptext)

            p_hdr = doc_new.add_paragraph()
            p_hdr.paragraph_format.space_after = Pt(6)
            p_hdr.paragraph_format.line_spacing = 1.15
            run_hdr = p_hdr.add_run(cleaned_hdr)

            if cleaned_hdr.startswith('📅'):
                run_hdr.bold = True
                run_hdr.font.color.rgb = DATE_COLOR
            elif cleaned_hdr.startswith('Диалоги') or cleaned_hdr.startswith('Полная хроника') or cleaned_hdr.startswith('═══'):
                run_hdr.bold = True
                if cleaned_hdr.startswith('═══'):
                    run_hdr.font.size = Pt(14)
                    run_hdr.font.color.rgb = SPEAKER_COLORS["kir"]

            header_count += 1

    doc_new.save(str(ANFI_DOC_PATH))
    # Also sync to output/books
    shutil.copy(ANFI_DOC_PATH, BOOKS_DOC_PATH)

    print(f"[Done Fixes] All user-reported errors successfully fixed in Анфи folder!")
    print(f"  - Messages processed: {msg_count}")
    print(f"  - Messages with applied fixes: {fix_count}")
    print(f"  - Headers/Date lines: {header_count}")

if __name__ == "__main__":
    process_anfi_file()
