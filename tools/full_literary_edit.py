# -*- coding: utf-8 -*-
"""
Полная литературная правка Голосовые_сообщения_2026.docx
- Удаление машинного мусора (Субтитры сделал, Отправить.)
- Исправление STT-ошибок (бренды, омофоны, бессмысленные слова)
- Устранение дублирований слов (я я, не не, ну ну)
- Добавление терминальных знаков
"""
import re
import docx
from docx.shared import RGBColor

INPUT = r'E:\coding\работа с литературой\output\books\Голосовые_сообщения_2026.docx'

ANFI_COLOR = RGBColor(0x00, 0x00, 0xFF)
KIR_COLOR = RGBColor(0xFF, 0x00, 0x00)

# Машинный мусор — удалять полностью или заменять
MACHINE_GARBAGE_PATTERNS = [
    r'Субтитры сделал\.?',
    r'Отправить\.\s*Отправить\.\s*',
    r'Прочай еще,?',
    r'\bход gou\b',
    r'\bЕцируuts\b',
    r'ецируuts',
    r'\bс redhares\b',
    r'\bredhares\b',
    r'С redhares',
    r'Happy end\. С redhares\.',
    r'\bHappy end\.?\s*',
    r'\bВальдерсе\b',
]

# Замены слов (контекстно-безопасные, только целые слова)
WORD_REPLACEMENTS = [
    # Бренды и программы
    (r'\bпротуз\b', 'Pro Tools'),
    (r'\bриппер\b', 'REAPER'),
    (r'\bриппере\b', 'REAPER'),
    (r'\bрепер\b', 'REAPER'),
    (r'\bкубейс\b', 'Cubase'),
    (r'\bевре перри\b', 'REAPER'),
    (r'\bевре\b', 'REAPER'),
    (r'\bf?lag player\b', 'FLAC-плеер'),
    (r'\bflag player\b', 'FLAC-плеер'),
    (r'\bизломанное\b', 'взломанное'),
    (r'\bмучим\b', 'там'),
    # STT-омофоны
    (r'\bвалясь\b', 'валяюсь'),
    (r'\bдва две\b', 'две'),
    (r'\bтри линиями\b', 'тремя линиями'),
    (r'\bбредят бредовый\b', 'бредовый'),
    (r'\bчеловечку зеленых\b', 'человечков в зелёном'),
    (r'\bбез вывозного\b', 'безвылазно'),
    (r'\bвэ этом\b', 'в этом'),
    (r'\bспесь специальные\b', 'специальные'),
    (r'\bспесь\b', 'специальные'),
    (r'\bшею же\b', 'это же'),
    (r'\bнож частично\b', 'ножки частично'),
    (r'\bнавесят шкаф\b', 'навесной шкаф'),
    (r'\bнавесят\b', 'навесной'),
    (r'\bджи машинка\b', 'LG-машинка'),
    (r'\bтуда-додвигать\b', 'туда-сюда двигать'),
    (r'\bтанковые люди\b', 'такие люди'),
    (r'\bпахеризм\b', 'пофигизм'),
    (r'\bбрюки не врубают\b', 'телефон не врубает'),
    (r'\bвашу мать\b', 'вашу мать'),
    (r'\bгринда\b', 'грит'),
    (r'\bлишь предпрямо\b', 'прямо'),
    (r'\bсуббота я такой\b', 'суббота'),
    (r'\bраджу руку\b', 'делал руку'),
    (r'\bдастин стал\b', 'Dustin Stahl'),
    (r'\bсins стал\b', 'Synth Stahl'),
    (r'\bна шокой\b', 'на такой'),
    (r'\bкинезащим сонить\b', 'кинематограф'),
    (r'\bтебя же\b', 'у тебя'),
    (r'\bвайд\b', 'wide'),
    (r'\bвтихаря\b', 'втихаря'),
    (r'\bмудилка-папашка\b', 'мудила-папашка'),
    (r'\bВальдбрис\b', 'Wildberries'),
    (r'\bВальдерсе\b', 'Wildberries'),
    (r'\bВБ банка\b', 'WB-банка'),
    (r'\bВБ\b', 'WB'),
    (r'\bснейр\b', 'snare'),
    (r'\bСнейр\b', 'snare'),
    (r'\bСнейра\b', 'snare'),
    (r'\bСУС\b', 'SUS'),
    (r'\bпит-шифтер\b', 'pitch-shifter'),
    (r'\bаудо-сити\b', 'Audacity'),
]

# Удаление дублирований
def fix_repetitions(text):
    text = re.sub(r'\b(я)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(не)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(ну)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(короче)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(вот)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(да)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(там)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(типа)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(прикольно)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    return text

# Терминальные знаки — заголовки новостей, ссылки, перечни
def fix_terminal_signs(content):
    # Пропускаем ссылки
    if content.startswith('Http') or content.startswith('https'):
        if not content.endswith(('.', '!', '?', '…')):
            content += '.'
        return content
    
    # Пропускаем короткие реплики и эмодзи
    if len(content) < 4:
        return content
    if content in ('🍿', '🍿.', 'Не.', 'Гемора куча смысла нету.', 'Эта про.', 'ReBirth'):
        if not content.endswith(('.', '!', '?', '…')):
            content += '.'
        return content
    
    # Новостные заголовки (содержат типичные слова)
    news_markers = ['РКН получил', 'ПРОВЕРЕНО:', 'В России хотят', 'Россия заняла',
                    'Россия проиграет', 'При оформлении SIM', 'ИИ-юрист',
                    'В России собираются', 'SimpleX Chat', 'Вежливость к ChatGPT',
                    'Скачайте лупы', 'HUAWEI AI Life', 'Скачать Wavelet']
    for marker in news_markers:
        if content.startswith(marker) or marker in content:
            if not content.endswith(('.', '!', '?', '…', ':', '—', '-', '»', ')', '|')):
                content += '.'
            return content
    
    # Прочие неполные
    if not content.endswith(('.', '!', '?', '…', ':', '—', '-', '»', ')', '(', '|', 
                              '!', ':)', ')))', '😊', '🌹', ' ')):
        if len(content) > 20 and not content.endswith(('))', ':)', '😀', '😂', '😎')):
            content += '.'
    return content


def apply_corrections(content):
    # 1. Удалить машинный мусор
    for pattern in MACHINE_GARBAGE_PATTERNS:
        content = re.sub(pattern, '', content, flags=re.IGNORECASE)
    
    # 2. Замены слов
    for pattern, replacement in WORD_REPLACEMENTS:
        content = re.sub(pattern, replacement, content, flags=re.IGNORECASE)
    
    # 3. Дублирования
    content = fix_repetitions(content)
    
    # 4. Лишние пробелы
    content = re.sub(r'\s+', ' ', content).strip()
    
    # 5. Терминальные знаки
    content = fix_terminal_signs(content)
    
    return content


def main():
    doc = docx.Document(INPUT)
    
    fixed_count = 0
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue
        
        # Проверяем, сообщение ли это
        msg_match = re.match(r'^(Анфи|Kir)\s+\[(\d{2}:\d{2})\]\s+\[([^\]]+)\]:\s*(.*)$', text)
        if not msg_match:
            # Обрабатываем даты и заголовки
            new_text = apply_corrections(text)
            if new_text != text:
                # Очищаем и пересоздаём
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
                fixed_count += 1
            continue
        
        speaker, time, msg_type, content = msg_match.groups()
        original_content = content
        
        # Применяем правки к содержимому
        new_content = apply_corrections(content)
        
        # Пересоздаём абзац с цветными именами
        if new_content != original_content or True:
            # Очищаем
            for run in para.runs:
                run.text = ''
            
            header = f"{speaker} [{time}] [{msg_type}]: "
            color = ANFI_COLOR if speaker == 'Анфи' else KIR_COLOR
            
            if para.runs:
                run = para.runs[0]
                run.text = header
                run.font.color.rgb = color
            else:
                run = para.add_run(header)
                run.font.color.rgb = color
            
            run2 = para.add_run(new_content)
            
            if new_content != original_content:
                fixed_count += 1
    
    doc.save(INPUT)
    print(f"Исправлено абзацев: {fixed_count}")
    print(f"Сохранено: {INPUT}")


if __name__ == '__main__':
    main()