#!/usr/bin/env python3
"""
Fixes "чай из этого получилось" -> "что из этого получилось" in:
E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx
"""

import sys
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

ANFI_DOC_PATH = Path(r"E:\coding\работа с литературой\Анфи\Диалоги_Анфи_и_Kir_2024-2026.docx")
BOOKS_DOC_PATH = Path(r"E:\coding\работа с литературой\output\books\Диалоги_Анфи_и_Kir_2024-2026.docx")

SPEAKER_COLORS = {
    "kir": RGBColor(13, 71, 161),
    "анфи": RGBColor(194, 24, 91),
    "default": RGBColor(74, 20, 140)
}
DATE_COLOR = RGBColor(46, 125, 50)

doc = Document(ANFI_DOC_PATH)
doc_new = Document()
style = doc_new.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

MSG_RE = re.compile(r'^(Kir|Анфи)\s+\[(\d{1,2}:\d{2})\]\s+\[(Голосовое|Текст)\]:\s*(.*)$', re.DOTALL)

for p in doc.paragraphs:
    ptext = p.text.strip()
    if not ptext:
        continue

    m = MSG_RE.match(ptext)
    if m:
        speaker = m.group(1)
        tstamp = m.group(2)
        mtype = m.group(3)
        body = m.group(4)

        body = re.sub(r'\bчай из этого получилось\b', 'что из этого получилось', body, flags=re.IGNORECASE)
        body = re.sub(r'\bчай из-за этого получилось\b', 'что из этого получилось', body, flags=re.IGNORECASE)

        p_new = doc_new.add_paragraph()
        p_new.paragraph_format.space_after = Pt(4)
        p_new.paragraph_format.line_spacing = 1.15

        header_prefix = f"{speaker} [{tstamp}] [{mtype}]: "
        run_spk = p_new.add_run(header_prefix)
        run_spk.bold = True
        run_spk.font.color.rgb = SPEAKER_COLORS.get(speaker.lower(), SPEAKER_COLORS["default"])

        p_new.add_run(body)
    else:
        p_hdr = doc_new.add_paragraph()
        p_hdr.paragraph_format.space_after = Pt(6)
        p_hdr.paragraph_format.line_spacing = 1.15
        run_hdr = p_hdr.add_run(ptext)

        if ptext.startswith('📅'):
            run_hdr.bold = True
            run_hdr.font.color.rgb = DATE_COLOR
        elif ptext.startswith('Диалоги') or ptext.startswith('Полная хроника') or ptext.startswith('═══'):
            run_hdr.bold = True
            if ptext.startswith('═══'):
                run_hdr.font.size = Pt(14)
                run_hdr.font.color.rgb = SPEAKER_COLORS["kir"]

doc_new.save(str(ANFI_DOC_PATH))
shutil.copy(ANFI_DOC_PATH, BOOKS_DOC_PATH)
print("Done fixing 'чай из этого получилось' -> 'что из этого получилось'.")
